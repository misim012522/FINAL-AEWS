from datetime import datetime, timezone
from bson import ObjectId
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Form
import csv
import re
import io
import unicodedata
try:
    import openpyxl
    _HAS_OPENPYXL = True
except Exception:
    _HAS_OPENPYXL = False
try:
    from docx import Document
    _HAS_PYTHON_DOCX = True
except Exception:
    _HAS_PYTHON_DOCX = False
from pymongo.errors import ServerSelectionTimeoutError
import os
from pathlib import Path
import shutil
from typing import Any, List
import uuid
from fastapi.responses import JSONResponse

from app.activity_log_utils import create_activity_log
from app.authz import get_current_actor
from app.database import get_db
from app.ai_features import build_model_feature_dict
from app.ai_model import _extract_topic_difficulty, _format_activity_title, predict_student_risk
from app.notification_utils import create_notification
from app.schemas import (
    BatchAddStudentsRequest,
    ClassCreate,
    ClassResponse,
    UpdateEnrollmentRequest,
)

router = APIRouter()

UPLOAD_DIR = Path(__file__).resolve().parent.parent.parent / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def _ensure_instructor_scope(actor: dict, instructor_id: str):
    if actor["role"] == "admin":
        return
    if actor["role"] != "instructor" or actor["id"] != (instructor_id or "").strip():
        raise HTTPException(status_code=403, detail="Forbidden")


def _get_class_for_actor(db, class_id: str, actor: dict):
    if not ObjectId.is_valid(class_id):
        raise HTTPException(status_code=404, detail="Class not found")
    doc = db.classes.find_one({"_id": ObjectId(class_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Class not found")
    _ensure_instructor_scope(actor, doc.get("instructor_id", ""))
    return doc


def _to_reason_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    normalized = _normalize_cell(value).strip().lower()
    return normalized in {"1", "true", "yes", "y", "checked"}


def _matches_numeric_referral_threshold(value, *, max_value: float) -> bool:
    if not isinstance(value, (int, float)):
        return False
    numeric = float(value)
    if numeric <= 0:
        return False
    return numeric <= max_value


def _matches_midterm_referral_threshold(value) -> bool:
    if not isinstance(value, (int, float)):
        return False
    numeric = float(value)
    if numeric <= 0:
        return False
    # Grade-scale midterms use 1.00-5.00.
    # This system treats 2.50 or below as the referral trigger.
    if numeric <= 5:
        return numeric <= 2.5
    # Percentage-style midterms use 0-100 where 75 or below is the trigger.
    return numeric <= 75


def _build_auto_referral_reasons(enrollment: dict) -> dict[str, bool]:
    midterm_grade = enrollment.get("midterm_grade")

    reasons = {
        "on_probation_status": False,
        "grade_2_5_or_below": _matches_midterm_referral_threshold(midterm_grade),
        "gwa_2_5_or_below": False,
        "low_midterm_performance": _to_reason_bool(enrollment.get("low_midterm_academic_performance")),
        "difficulty_catching_up": False,
    }
    return reasons


def _find_amu_staff_for_college(db, college: str) -> dict | None:
    college_name = _normalize_cell(college).strip()
    if not college_name:
        return None
    return db.amustaff.find_one(
        {
            "archived": {"$ne": True},
            "status": {"$in": ["active", None, ""]},
            "college": {"$regex": f"^{re.escape(college_name)}$", "$options": "i"},
        },
        sort=[("name", 1), ("_id", 1)],
    )


def _apply_automatic_referral(db, class_doc: dict, enrollment_doc: dict) -> bool:
    reasons = _build_auto_referral_reasons(enrollment_doc)
    existing_reasons = enrollment_doc.get("referral_reasons") or {}
    already_referred = enrollment_doc.get("flagged_for_mentoring") is True

    if not any(reasons.values()):
        if already_referred and any(existing_reasons.values()):
            db.enrollments.update_one(
                {"_id": enrollment_doc["_id"]},
                {
                    "$set": {"flagged_for_mentoring": False},
                    "$unset": {
                        "referral_reasons": "",
                        "assigned_amu_staff_id": "",
                        "assigned_amu_staff_name": "",
                        "assigned_amu_staff_college": "",
                        "referred_at": "",
                    },
                },
            )
            return True
        return False

    instructor_doc = None
    instructor_id = class_doc.get("instructor_id")
    if instructor_id and ObjectId.is_valid(instructor_id):
        instructor_doc = db.instructor.find_one({"_id": ObjectId(instructor_id)})
    instructor_college = _normalize_cell((instructor_doc or {}).get("college"))

    assigned_staff = _find_amu_staff_for_college(db, instructor_college)
    if not assigned_staff:
        return False

    assigned_staff_id = str(assigned_staff["_id"])
    assigned_staff_name = _normalize_cell(assigned_staff.get("name"))
    assigned_staff_college = _normalize_cell(assigned_staff.get("college")) or None
    already_assigned = (
        _normalize_cell(enrollment_doc.get("assigned_amu_staff_id")) == assigned_staff_id
        and _normalize_cell(enrollment_doc.get("assigned_amu_staff_name")) == assigned_staff_name
        and _normalize_cell(enrollment_doc.get("assigned_amu_staff_college")) == (assigned_staff_college or "")
    )
    same_reasons = existing_reasons == reasons
    if already_referred and already_assigned and same_reasons:
        return False

    update_data = {
        "flagged_for_mentoring": True,
        "referral_reasons": reasons,
        "assigned_amu_staff_id": assigned_staff_id,
        "assigned_amu_staff_name": assigned_staff_name,
        "assigned_amu_staff_college": assigned_staff_college,
    }
    if not enrollment_doc.get("referred_at"):
        update_data["referred_at"] = datetime.now(timezone.utc)

    db.enrollments.update_one({"_id": enrollment_doc["_id"]}, {"$set": update_data})

    subject_code = class_doc.get("subject_code") or "Class"
    subject_name = class_doc.get("subject_name") or ""
    class_label = f"{subject_code}" + (f": {subject_name}" if subject_name else "")
    student_label = (
        _normalize_cell(enrollment_doc.get("student_email"))
        or _normalize_cell(enrollment_doc.get("student_id"))
        or _normalize_cell(enrollment_doc.get("student_name"))
        or "Student"
    )
    instructor_name = _normalize_cell((instructor_doc or {}).get("name")) or "Instructor"

    if not already_referred:
        create_notification(
            db,
            role="amu-staff",
            recipient_user_id=assigned_staff_id,
            title="Student auto-referred by system",
            body=(
                f"Student {student_label} met the automatic referral criteria in "
                f"{class_label} and was assigned to {assigned_staff_name}"
                f"{f' - {assigned_staff_college}' if assigned_staff_college else ''}. "
                f"Please review the case."
            ),
            type="alert",
        )
        create_notification(
            db,
            role="instructor",
            recipient_user_id=str(class_doc.get("instructor_id")) if class_doc and class_doc.get("instructor_id") else None,
            title="Student auto-referred to AMU",
            body=(
                f"{student_label} from {class_label} was automatically referred to "
                f"{assigned_staff_name}{f' - {assigned_staff_college}' if assigned_staff_college else ''} "
                f"based on the recorded referral reasons."
            ),
            type="info",
        )

    return True


def _set_related_enrollment_archive_state(db, class_id: str, archived: bool) -> None:
    """Keep enrollment records in sync with their parent class archive state."""
    if archived:
        db.enrollments.update_many(
            {"class_id": class_id},
            {
                "$set": {
                    "archived": True,
                    "status": "archived",
                    "archived_at": datetime.now(timezone.utc),
                }
            },
        )
        return

    db.enrollments.update_many(
        {"class_id": class_id},
        {
            "$set": {"status": "active"},
            "$unset": {"archived": "", "archived_at": ""},
        },
    )

_HEADER_HINTS = [
    "email", "name", "student", "id", "no", "number",
    "grade", "score", "midterm", "final", "lab", "standing",
    "attendance", "date", "signature", "present", "absent",
    "section", "subject", "course", "code",
]

_SUBJECT_CODE_PATTERN = re.compile(
    r"(?:subject|course)\s*code\s*[:\-]?\s*([A-Za-z0-9][A-Za-z0-9_./\- ]{1,40})",
    re.IGNORECASE,
)


def _normalize_cell(value) -> str:
    if value is None:
        return ""
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        # Excel IDs often appear as 2201103564.0; normalize to integer text.
        if value.is_integer():
            return str(int(value))
        return str(value).strip()

    text = str(value).strip()
    if re.fullmatch(r"\d+\.0", text):
        return text[:-2]
    return text


def _make_unique_headers(header_cells) -> list[str]:
    headers = []
    seen = {}
    for i, cell in enumerate(header_cells):
        base = _normalize_cell(cell).lower() or f"col_{i}"
        count = seen.get(base, 0)
        headers.append(base if count == 0 else f"{base}_{count + 1}")
        seen[base] = count + 1
    return headers


def _slugify_header_fragment(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", _normalize_cell(value))
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = ascii_text.lower().strip()
    ascii_text = re.sub(r"[^a-z0-9]+", "_", ascii_text)
    return ascii_text.strip("_")


def _header_term_to_prefix(value: str) -> str | None:
    normalized = _normalize_cell(value).lower().strip()
    if not normalized:
        return None
    if normalized in {"final", "finals", "final term"}:
        return "finals"
    if normalized in {"mid", "midterm", "midterms", "mid term"}:
        return "midterm"
    return None


def _normalize_component_bucket_label(value: str) -> str | None:
    normalized = _normalize_cell(value).lower().strip()
    if not normalized:
        return None
    if "class standing" in normalized or normalized in {"cs", "class stand"}:
        return "class_standing"
    if "laboratory" in normalized or normalized in {"lab", "laboratory"}:
        return "laboratory"
    if "major output" in normalized or normalized in {"mo", "major"}:
        return "major_output"
    return None


def _build_gradesheet_multirow_headers(raw_rows) -> tuple[list[str], int] | tuple[None, None]:
    """Build canonical headers from multi-row gradesheet layouts with section/title rows."""
    if not raw_rows:
        return None, None

    max_scan = min(len(raw_rows), 12)
    column_count = max((len(row) for row in raw_rows[:max_scan]), default=0)
    if column_count <= 0:
        return None, None

    term_by_col: dict[int, str] = {}
    bucket_by_col: dict[int, str] = {}
    title_by_col: dict[int, str] = {}
    student_header_row = None
    student_id_col = None
    student_name_col = None
    last_header_row = 0

    for row_idx, row in enumerate(raw_rows[:max_scan]):
        if (
            student_header_row is not None
            and student_id_col is not None
            and student_name_col is not None
            and row_idx > student_header_row
        ):
            candidate_id = _normalize_cell(row[student_id_col] if student_id_col < len(row) else "").strip()
            candidate_name = _normalize_cell(row[student_name_col] if student_name_col < len(row) else "").strip()
            if candidate_id and candidate_name:
                break

        current_term = None
        current_bucket = None
        for col_idx in range(column_count):
            raw_value = row[col_idx] if col_idx < len(row) else None
            text = _normalize_cell(raw_value).strip()
            lower = text.lower()

            term = _header_term_to_prefix(lower) if lower else None
            if term:
                current_term = term
                last_header_row = max(last_header_row, row_idx)
            if current_term:
                term_by_col[col_idx] = current_term

            bucket = _normalize_component_bucket_label(lower) if lower else None
            if bucket:
                current_bucket = bucket
                last_header_row = max(last_header_row, row_idx)
            if current_bucket:
                bucket_by_col[col_idx] = current_bucket

            if not text:
                continue

            if lower in {"student id", "student_id", "id number"}:
                student_header_row = row_idx
                student_id_col = col_idx
                last_header_row = max(last_header_row, row_idx)
            elif lower in {"student name", "name of students", "name of student", "name"}:
                student_header_row = row_idx if student_header_row is None else student_header_row
                student_name_col = col_idx
                last_header_row = max(last_header_row, row_idx)

            if row_idx == 0:
                continue
            if lower in {"gpe", "gpa", "equi", "equivalent", "mtg", "ftg", "fg", "summary"}:
                continue
            if _header_term_to_prefix(lower):
                continue
            if _normalize_component_bucket_label(lower):
                continue
            if re.fullmatch(r"\d+", lower):
                continue

            slug = _slugify_header_fragment(text)
            if slug:
                title_by_col[col_idx] = slug
                last_header_row = max(last_header_row, row_idx)

    if student_header_row is None or student_id_col is None or student_name_col is None:
        return None, None

    headers: list[str] = []
    for col_idx in range(column_count):
        if col_idx == student_id_col:
            headers.append("student_id")
            continue
        if col_idx == student_name_col:
            headers.append("student_name")
            continue

        term = term_by_col.get(col_idx)
        bucket = bucket_by_col.get(col_idx)
        title = title_by_col.get(col_idx)
        if term and bucket and title:
            headers.append(f"{term}_{bucket}_{title}")
            continue

        source_value = raw_rows[student_header_row][col_idx] if col_idx < len(raw_rows[student_header_row]) else None
        headers.append(_normalize_cell(source_value).lower() or f"col_{col_idx}")

    if sum(1 for header in headers if header.startswith(("midterm_", "finals_"))) < 6:
        return None, None

    return _make_unique_headers(headers), last_header_row


def _guess_header_index(raw_rows) -> int | None:
    best_idx = None
    best_score = 0

    for idx, row in enumerate(raw_rows[:30]):
        normalized_cells = [_normalize_cell(cell).lower() for cell in row if _normalize_cell(cell)]
        if len(normalized_cells) < 2:
            continue

        score = sum(1 for cell in normalized_cells if any(hint in cell for hint in _HEADER_HINTS))
        if score > best_score:
            best_idx = idx
            best_score = score

    return best_idx if best_score > 0 else None


def _rows_to_dicts(raw_rows) -> list[dict]:
    rows = []
    if not raw_rows:
        return rows

    multi_headers, data_start_header_idx = _build_gradesheet_multirow_headers(raw_rows)
    if multi_headers is not None and data_start_header_idx is not None:
        for r in raw_rows[data_start_header_idx + 1:]:
            row_dict = {}
            for i, val in enumerate(r):
                key = multi_headers[i] if i < len(multi_headers) else f"col_{i}"
                row_dict[key] = _normalize_cell(val)
            if any(v for v in row_dict.values()):
                rows.append(row_dict)
        if rows:
            return rows

    header_idx = _guess_header_index(raw_rows)
    if header_idx is None:
        for r in raw_rows:
            row_dict = {f"col_{i}": _normalize_cell(v) for i, v in enumerate(r)}
            if any(v for v in row_dict.values()):
                rows.append(row_dict)
        return rows

    headers = _make_unique_headers(raw_rows[header_idx])
    for r in raw_rows[header_idx + 1:]:
        row_dict = {}
        for i, val in enumerate(r):
            key = headers[i] if i < len(headers) else f"col_{i}"
            row_dict[key] = _normalize_cell(val)
        if any(v for v in row_dict.values()):
            rows.append(row_dict)
    return rows


def _resolve_workbook_term_prefix(sheet_name: str) -> str | None:
    normalized = _normalize_cell(sheet_name).lower()
    if "midterm" in normalized or "mid term" in normalized:
        return "midterm"
    if "finalterm" in normalized or "final term" in normalized or "finals" in normalized or "final" in normalized:
        return "finals"
    return None


def _extract_buksu_raw_activity_rows(ws, term_prefix: str) -> list[dict]:
    """Parse BukSU raw-score worksheets where activity titles live on rows 15-17."""
    if ws.max_row < 18 or ws.max_column < 8:
        return []

    id_header = _normalize_cell(ws.cell(row=13, column=6).value).lower().strip()
    name_header = _normalize_cell(ws.cell(row=13, column=7).value).lower().strip()
    if id_header not in {"id number", "student id", "student no."} or name_header not in {"name", "name of students"}:
        return []

    headers_by_col: dict[int, str] = {}
    current_bucket = None

    for col_idx in range(8, ws.max_column + 1):
        bucket_value = _normalize_cell(ws.cell(row=13, column=col_idx).value).lower().strip()
        detected_bucket = _normalize_component_bucket_label(bucket_value)
        if detected_bucket:
            current_bucket = detected_bucket

        if not current_bucket:
            continue

        title = _normalize_cell(ws.cell(row=16, column=col_idx).value).strip()
        if not title:
            continue

        slug = _slugify_header_fragment(title)
        if not slug:
            continue

        headers_by_col[col_idx] = f"{term_prefix}_{current_bucket}_{slug}"

    if not headers_by_col:
        return []

    parsed_rows: list[dict] = []
    for row_idx in range(18, ws.max_row + 1):
        student_id = _normalize_cell(ws.cell(row=row_idx, column=6).value).strip()
        student_name = _normalize_cell(ws.cell(row=row_idx, column=7).value).strip()

        if not student_id and not student_name:
            continue

        row_dict = {
            "student_id": student_id,
            "student_name": student_name,
        }

        has_scores = False
        for col_idx, header in headers_by_col.items():
            value = _normalize_cell(ws.cell(row=row_idx, column=col_idx).value).strip()
            if not value:
                continue
            row_dict[header] = value
            has_scores = True

        if has_scores or student_id or student_name:
            parsed_rows.append(row_dict)

    return parsed_rows


def _merge_gradesheet_row_sets(base_rows: list[dict], extra_rows: list[dict]) -> list[dict]:
    """Merge workbook row sets by student id/name so summary and raw activity sheets combine."""
    if not base_rows:
        return extra_rows
    if not extra_rows:
        return base_rows

    def build_keys(row: dict) -> list[str]:
        keys: list[str] = []
        student_id = _normalize_cell(row.get("student_id") or row.get("id number") or row.get("student no.")).strip().lower()
        student_name = _normalize_cell(row.get("student_name") or row.get("name of students") or row.get("name")).strip().lower()
        if student_id:
            keys.append(f"id:{student_id}")
        if student_name:
            keys.append(f"name:{student_name}")
        return keys

    merged_rows = [dict(row) for row in base_rows]
    row_lookup: dict[str, dict] = {}
    for row in merged_rows:
        for key in build_keys(row):
            row_lookup.setdefault(key, row)

    for extra_row in extra_rows:
        target = None
        for key in build_keys(extra_row):
            if key in row_lookup:
                target = row_lookup[key]
                break
        if target is None:
            target = dict(extra_row)
            merged_rows.append(target)
        else:
            for key, value in extra_row.items():
                if value not in (None, ""):
                    target[key] = value

        for key in build_keys(target):
            row_lookup[key] = target

    return merged_rows


def _score_rows_for_preferred_type(rows: list[dict], preferred_type: str | None) -> int:
    if not rows:
        return -1
    if not preferred_type:
        return 0

    keys = list(rows[0].keys())
    score = 0

    id_col = _find_student_id_column(keys)
    name_col = _find_column(keys, ['name of students', 'student name', 'name'])
    if id_col:
        score += 10
    if name_col:
        score += 10

    if preferred_type == "gradesheet":
        for kw in ['mtg', 'ftg', 'fg', 'midterm', 'final', 'class standing', 'cs (30%)', 'lab (30%)', 'mo (40%)']:
            if _find_column(keys, [kw]):
                score += 8
        if _find_column(keys, ['mtg(1/3)']) and _find_column(keys, ['ftg(2/3)']):
            score += 20
    elif preferred_type == "attendance":
        for kw in ['attendance', 'present', 'absent', 'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']:
            if _find_column(keys, [kw]):
                score += 6
    elif preferred_type == "classlist":
        if id_col:
            score += 20
        if name_col:
            score += 20

    return score


def _to_number_or_text(value: str):
    text = _normalize_cell(value)
    if not text:
        return None

    clean = text.replace('%', '').strip()
    try:
        return float(clean)
    except (TypeError, ValueError):
        return text


def _to_bool(value) -> bool | None:
    text = _normalize_cell(value).strip().lower()
    if not text:
        return None
    if text in {"1", "true", "yes", "y", "present", "checked"}:
        return True
    if text in {"0", "false", "no", "n", "absent", "unchecked"}:
        return False
    return None


def _to_float(value) -> float | None:
    parsed = _to_number_or_text(_normalize_cell(value))
    if isinstance(parsed, (int, float)):
        return float(parsed)
    return None


def _to_int(value) -> int | None:
    parsed = _to_float(value)
    if parsed is None:
        return None
    return int(parsed)


_AI_UPLOAD_FIELD_MAPPINGS = {
    "previous_gpa": ["previous gpa", "previous_gpa", "gpa"],
    "failed_subject_count": ["failed subject count", "failed_subject_count", "failed subjects", "failed_subjects"],
    "self_reported_attendance": ["attendance", "attendance rate", "attendance_rate"],
    "received_academic_support": ["received academic support", "received_academic_support", "academic support"],
    "difficulty_understanding_lectures": ["difficulty in understanding lectures", "difficulty_understanding_lectures"],
    "struggles_specific_subjects": ["struggles with specific subjects", "struggles_specific_subjects"],
    "weak_study_habits_time_management": ["weak study habits or time management", "weak_study_habits_time_management"],
    "low_motivation_engagement": ["low motivation or engagement", "low_motivation_engagement"],
    "poor_comprehension_writing_skills": ["poor comprehension or writing skills", "poor_comprehension_writing_skills"],
    "financial_difficulties": ["financial difficulties", "financial_difficulties"],
    "physical_health_concerns": ["physical health-related concerns", "physical_health_concerns"],
    "family_issues": ["family issues", "family_issues"],
    "part_time_work_affecting_studies": ["part-time work affecting studies", "part_time_work_affecting_studies"],
    "mental_health_concerns": ["mental health-related concerns", "mental_health_concerns"],
    "internet_issues": ["internet/connectivity issues", "internet connectivity issues", "internet issues", "connectivity issues", "internet_issues", "connectivity_issues"],
}

_PREVIOUS_GRADES_FIELD_MAPPINGS = {
    "previous_gpa": ["previous gpa", "previous_gpa", "gpa"],
    "previous_midterm_class_standing": ["previous midterm class standing", "previous_midterm_class_standing", "cs (30%)"],
    "previous_midterm_laboratory": ["previous midterm laboratory", "previous_midterm_laboratory", "lab (30%)", "lab (40%)"],
    "previous_midterm_major_output": ["previous midterm major output", "previous_midterm_major_output", "mo (40%)"],
    "previous_midterm_grade": ["previous midterm grade", "previous_midterm_grade", "previous mtg", "previous_mtg", "mtg", "mtg(1/3)"],
    "previous_final_class_standing": ["previous final class standing", "previous_final_class_standing", "cs (30%)_2"],
    "previous_final_laboratory": ["previous final laboratory", "previous_final_laboratory", "lab (30%)_2"],
    "previous_final_major_output": ["previous final major output", "previous_final_major_output", "mo (40%)"],
    "previous_final_grade": ["previous final grade", "previous_final_grade", "previous ftg", "previous_ftg", "previous fg", "fg", "ftg", "ftg(2/3)"],
    "previous_failed_flag": ["previous failed flag", "previous_failed_flag", "previously failed", "has previous failure"],
    "previous_passed_flag": ["previous passed flag", "previous_passed_flag", "previously passed", "has previous pass"],
    "historical_grade_average": ["historical grade average", "historical_grade_average", "grade history average", "average previous grade", "fg", "ftg", "ftg(2/3)"],
    "historical_failure_count": ["historical failure count", "historical_failure_count", "historical failed subjects", "previous failures", "failure history count"],
}


def _build_ai_input_update_data(row: dict, keys: list[str]) -> dict:
    update_data = {}
    numeric_fields = {"previous_gpa", "self_reported_attendance"}
    integer_fields = {"failed_subject_count"}
    boolean_fields = set(_AI_UPLOAD_FIELD_MAPPINGS.keys()) - numeric_fields - integer_fields

    for db_field, aliases in _AI_UPLOAD_FIELD_MAPPINGS.items():
        col_name = _find_column(keys, aliases)
        if not col_name:
            continue

        raw_value = row.get(col_name, "")
        if db_field in numeric_fields:
            parsed = _to_float(raw_value)
        elif db_field in integer_fields:
            parsed = _to_int(raw_value)
        elif db_field in boolean_fields:
            parsed = _to_bool(raw_value)
        else:
            parsed = raw_value

        if parsed is None or parsed == "":
            continue
        update_data[db_field] = parsed

    return update_data


def _build_previous_grades_update_data(row: dict, keys: list[str]) -> dict:
    update_data = {}
    numeric_fields = {
        "previous_gpa",
        "previous_midterm_class_standing",
        "previous_midterm_laboratory",
        "previous_midterm_major_output",
        "previous_midterm_grade",
        "previous_final_class_standing",
        "previous_final_laboratory",
        "previous_final_major_output",
        "previous_final_grade",
        "historical_grade_average",
    }
    integer_fields = {
        "historical_failure_count",
        "previous_failed_flag",
        "previous_passed_flag",
    }

    for db_field, aliases in _PREVIOUS_GRADES_FIELD_MAPPINGS.items():
        col_name = _find_column(keys, aliases)
        if not col_name:
            continue

        raw_value = row.get(col_name, "")
        if db_field in numeric_fields:
            parsed = _to_float(raw_value)
        elif db_field in integer_fields:
            parsed = _to_int(raw_value)
        else:
            parsed = raw_value

        if parsed is None or parsed == "":
            continue
        update_data[db_field] = parsed

    remarks_col = _find_column(keys, ["remarks", "remark", "status"])
    remarks_value = _normalize_cell(row.get(remarks_col, "")) if remarks_col else ""
    remarks_normalized = remarks_value.strip().lower()
    if remarks_normalized:
        if "fail" in remarks_normalized:
            update_data.setdefault("previous_failed_flag", 1)
            update_data.setdefault("previous_passed_flag", 0)
            update_data.setdefault("historical_failure_count", 1)
        elif "pass" in remarks_normalized:
            update_data.setdefault("previous_failed_flag", 0)
            update_data.setdefault("previous_passed_flag", 1)
            update_data.setdefault("historical_failure_count", 0)

    previous_failed_flag = update_data.get("previous_failed_flag")
    previous_passed_flag = update_data.get("previous_passed_flag")
    historical_failure_count = update_data.get("historical_failure_count")

    if historical_failure_count is not None:
        if previous_failed_flag is None:
            update_data["previous_failed_flag"] = 1 if int(historical_failure_count) > 0 else 0
        if previous_passed_flag is None:
            update_data["previous_passed_flag"] = 0 if int(historical_failure_count) > 0 else 1
    elif previous_failed_flag is not None and previous_passed_flag is None:
        update_data["previous_passed_flag"] = 0 if int(previous_failed_flag) else 1
    elif previous_passed_flag is not None and previous_failed_flag is None:
        update_data["previous_failed_flag"] = 0 if int(previous_passed_flag) else 1

    if "historical_grade_average" not in update_data:
        if update_data.get("previous_gpa") is not None:
            update_data["historical_grade_average"] = float(update_data["previous_gpa"])
        elif update_data.get("previous_final_grade") is not None:
            update_data["historical_grade_average"] = float(update_data["previous_final_grade"])

    if "previous_gpa" not in update_data:
        if update_data.get("previous_final_grade") is not None:
            update_data["previous_gpa"] = float(update_data["previous_final_grade"])
        elif update_data.get("previous_midterm_grade") is not None:
            update_data["previous_gpa"] = float(update_data["previous_midterm_grade"])

    return update_data


def _is_score_column(key: str) -> bool:
    if not key:
        return False
    normalized = key.lower().strip()
    grade_tokens = [
        'class standing', 'cs (30%)', 'lab', 'laboratory', 'mo (40%)', 'major output',
        'summary', 'midterm', 'mtg', 'final', 'ftg', 'fg', 'quiz', 'exam',
        'score', 'project', 'activity', 'attendance'
    ]
    if any(token in normalized for token in grade_tokens):
        return True

    # Support compact headers often found in sheets, e.g. CS, MO, LAB, FG, MTG, FTG.
    return any(
        re.fullmatch(pattern, normalized)
        for pattern in (
            r"cs(?:_\d+)?",
            r"mo(?:_\d+)?",
            r"lab(?:_\d+)?",
            r"mtg(?:_\d+)?",
            r"ftg(?:_\d+)?",
            r"fg(?:_\d+)?",
        )
    )


def _is_raw_component_score_column(key: str) -> bool:
    """True when a column is a raw CS/LAB/MO component score column."""
    if not key:
        return False

    normalized = key.lower().strip()

    if _is_grade_output_column(normalized):
        return False

    component_tokens = [
        'class standing', 'cs (30%)', 'cs',
        'lab (30%)', 'lab', 'laboratory',
        'mo (40%)', 'major output', 'mo',
        'class_standing', 'major_output',
    ]
    if any(token in normalized for token in component_tokens):
        return True

    return any(
        re.fullmatch(pattern, normalized)
        for pattern in (
            r"cs(?:_\d+)?",
            r"lab(?:_\d+)?",
            r"mo(?:_\d+)?",
            r"(?:midterm|finals)_(?:class_standing|laboratory|major_output)_.+",
        )
    )


def _is_grade_output_column(key: str) -> bool:
    """Detect computed/aggregate grade columns that should not appear in raw scores."""
    if not key:
        return False

    normalized = key.lower().strip()

    exact = {
        "fg", "ftg", "mtg", "gpa",
        "overall grade", "midterm grade", "final grade",
        "summary",
        "midterm_weighted", "final_weighted",
        "mtg(1/3)", "ftg(2/3)",
    }
    if normalized in exact:
        return True

    return any(
        token in normalized
        for token in (
            "gpa", "overall grade", "weighted",
            "midterm grade", "final grade", "finals grade",
            "final computed", "general average",
        )
    )


def _looks_like_raw_fraction_score(value: str) -> bool:
    """True for values like 10/20 or 17.5 / 25."""
    if not value:
        return False
    return re.fullmatch(r"\d+(?:\.\d+)?\s*/\s*\d+(?:\.\d+)?", value.strip()) is not None


def _detect_raw_component_bucket(key: str) -> str | None:
    """Detect CS/LAB/MO raw component section from a header key."""
    if not key:
        return None

    normalized = key.lower().strip()

    if (
        "class standing" in normalized
        or "class_standing" in normalized
        or "cs (30%)" in normalized
        or re.fullmatch(r"cs(?:_\d+)?", normalized)
    ):
        return "class standing"

    if (
        "laboratory" in normalized
        or "lab (30%)" in normalized
        or re.fullmatch(r"lab(?:_\d+)?", normalized)
        or re.search(r"\blab\b", normalized)
        or "midterm_laboratory_" in normalized
        or "finals_laboratory_" in normalized
    ):
        return "laboratory"

    if (
        "major output" in normalized
        or "major_output" in normalized
        or "mo (40%)" in normalized
        or re.fullmatch(r"mo(?:_\d+)?", normalized)
        or re.search(r"\bmo\b", normalized)
    ):
        return "major output"

    return None


def _build_raw_score_column_aliases(
    keys: list[str],
    identity_cols: set[str] | None = None,
    alias_all_section_columns: bool = True,
) -> dict[str, str]:
    """Build stable aliases for raw score columns inside CS/LAB/MO sections."""
    aliases = {}
    identity_cols = set(identity_cols or set())

    canonical_dynamic_keys = [
        col
        for col in keys
        if _normalize_cell(col).lower().strip().startswith(("midterm_", "finals_"))
    ]
    if canonical_dynamic_keys:
        for col in canonical_dynamic_keys:
            normalized = _normalize_cell(col).lower().strip()
            if normalized:
                aliases[col] = normalized
        return aliases

    current_bucket = None
    bucket_counts = {
        "class standing": 0,
        "laboratory": 0,
        "major output": 0,
    }

    base_section_headers = {
        "class standing", "cs (30%)",
        "laboratory", "lab (30%)",
        "major output", "mo (40%)",
    }

    skip_metadata = {
        "ctrl no", "control no", "no", "no.",
        "email", "student email",
        "name", "full name", "student name", "name of students", "name of student",
        "id", "id number", "student id", "student no", "student no.", "sid",
        "section", "section code", "sec",
        "subject", "subject code", "course", "course code",
        "class time", "time", "schedule", "sched",
    }

    for col in keys:
        if not col or col in identity_cols:
            continue

        normalized = _normalize_cell(col).lower().strip()
        if not normalized:
            continue

        if normalized.startswith(("midterm_", "finals_")):
            aliases[col] = normalized
            continue

        detected_bucket = _detect_raw_component_bucket(normalized)
        if detected_bucket:
            current_bucket = detected_bucket

        if _is_grade_output_column(normalized) or normalized in skip_metadata:
            continue

        if current_bucket is None:
            continue

        needs_alias = (
            alias_all_section_columns
            or normalized.startswith("col_")
            or normalized in base_section_headers
        )
        if not needs_alias:
            continue

        bucket_counts[current_bucket] += 1
        aliases[col] = f"{current_bucket} item {bucket_counts[current_bucket]}"

    return aliases


def _normalize_score_term(value: str | None) -> str | None:
    normalized = _normalize_cell(value).lower().strip()
    if not normalized:
        return None
    if normalized.startswith("mid"):
        return "midterm"
    if normalized.startswith("fin"):
        return "final"
    if normalized.startswith("finals"):
        return "final"
    return None


def _is_midterm_to_final_boundary_column(key: str) -> bool:
    normalized = _normalize_cell(key).lower().strip()
    return normalized in {
        "mtg", "midterm", "midterm grade", "mid", "mt", "mtg(1/3)", "midterm_weighted",
    }


def _is_final_term_hint_column(key: str) -> bool:
    normalized = _normalize_cell(key).lower().strip()
    if not normalized:
        return False

    if normalized in {"ftg", "fg", "final", "final grade", "finals", "fin", "ftg(2/3)", "final_weighted"}:
        return True

    if "final" in normalized:
        return True

    if "_2" in normalized and any(token in normalized for token in ("cs", "lab", "mo", "class standing", "laboratory", "major output")):
        return True

    return False


def _build_raw_score_column_terms(
    keys: list[str],
    score_aliases: dict[str, str],
    identity_cols: set[str] | None = None,
) -> dict[str, str]:
    """Tag raw score source columns as midterm/final based on header flow."""
    identity_cols = set(identity_cols or set())
    terms = {}
    current_term = "midterm"

    for col in keys:
        if not col or col in identity_cols:
            continue

        normalized = _normalize_cell(col).lower().strip()
        if not normalized:
            continue

        if normalized.startswith("midterm_"):
            terms[col] = "midterm"
            continue
        if normalized.startswith("finals_"):
            terms[col] = "final"
            continue

        if _is_midterm_to_final_boundary_column(normalized):
            current_term = "final"
            continue

        if _is_final_term_hint_column(normalized):
            current_term = "final"

        if col in score_aliases:
            terms[col] = current_term

    return terms


def _format_score_column_display_name(score_col: str, score_term: str | None = None) -> str:
    normalized_col = _normalize_cell(score_col).strip()
    pretty_col = normalized_col.title()
    normalized_term = _normalize_score_term(score_term)
    if normalized_term == "midterm":
        return f"Midterm - {pretty_col}"
    if normalized_term == "final":
        return f"Final - {pretty_col}"
    return pretty_col


def _collect_class_score_column_metadata(enrollments: list[dict]) -> tuple[list[str], dict[str, str]]:
    score_columns: list[str] = []
    score_column_terms: dict[str, str] = {}
    has_breakdown_scores = any(bool(enrollment.get("grades_breakdown")) for enrollment in enrollments)

    for enrollment in enrollments:
        stored_scores = dict(enrollment.get("grades_breakdown") or {})
        stored_column_order = list(enrollment.get("grades_column_order") or list(stored_scores.keys()))
        stored_terms = dict(enrollment.get("grades_column_terms") or {})

        for source_col in stored_column_order:
            if source_col not in stored_scores:
                continue
            value = stored_scores[source_col]
            if not _should_include_score_column(source_col, _normalize_cell(value), value):
                continue
            inferred_term = (
                _normalize_score_term(stored_terms.get(source_col))
                or ("final" if _is_final_term_hint_column(source_col) else None)
                or "midterm"
            )
            score_column_terms.setdefault(source_col, inferred_term)
            if source_col not in score_columns:
                score_columns.append(source_col)

        if has_breakdown_scores:
            continue

        legacy_scores = _legacy_grade_score_map(enrollment)
        legacy_aliases = _build_raw_score_column_aliases(
            list(legacy_scores.keys()),
            set(),
            alias_all_section_columns=True,
        )
        for legacy_col, legacy_value in legacy_scores.items():
            score_col = legacy_aliases.get(legacy_col, legacy_col)
            if not _should_include_score_column(score_col, _normalize_cell(legacy_value), legacy_value):
                continue
            score_column_terms.setdefault(score_col, "final" if legacy_col.startswith("final_") else "midterm")
            if score_col not in score_columns:
                score_columns.append(score_col)

    score_columns = sorted(score_columns, key=_raw_score_column_sort_key)
    score_column_terms = _infer_missing_score_column_terms(score_columns, score_column_terms)
    return score_columns, score_column_terms


def _build_score_column_lookup(score_columns: list[str], score_column_terms: dict[str, str]) -> dict[str, str]:
    lookup: dict[str, str] = {}

    for score_col in score_columns:
        normalized_col = _normalize_cell(score_col).lower().strip()
        if normalized_col:
            lookup.setdefault(normalized_col, score_col)

        display_name = _format_score_column_display_name(score_col, score_column_terms.get(score_col))
        normalized_display = _normalize_cell(display_name).lower().strip()
        if normalized_display:
            lookup.setdefault(normalized_display, score_col)

    return lookup


def _parse_activity_item_number(value: Any) -> int | None:
    text = _normalize_cell(value).strip()
    if not text:
        return None
    match = re.search(r"(\d+)", text)
    if not match:
        return None
    return int(match.group(1))


def _resolve_mapping_target_column(
    row: dict[str, Any],
    keys: list[str],
    score_columns: list[str],
    score_column_terms: dict[str, str],
    score_column_lookup: dict[str, str],
) -> str | None:
    column_key = _find_column(
        keys,
        [
            "activity column", "activity_column", "score column", "score_column",
            "column", "feature", "feature key", "feature_key", "activity key", "activity_key",
        ],
    )
    if column_key:
        normalized_value = _normalize_cell(row.get(column_key)).lower().strip()
        if normalized_value:
            matched = score_column_lookup.get(normalized_value)
            if matched:
                return matched

    term_key = _find_column(keys, ["term", "grading term", "period"])
    component_key = _find_column(keys, ["component", "bucket", "section", "category"])
    item_key = _find_column(keys, ["item", "item no", "item number", "activity no", "activity number", "order", "index"])

    term_value = _normalize_score_term(row.get(term_key)) if term_key else None
    component_value = _detect_raw_component_bucket(row.get(component_key)) if component_key else None
    item_number = _parse_activity_item_number(row.get(item_key)) if item_key else None

    if term_value and component_value and item_number:
        matching_columns = [
            score_col
            for score_col in score_columns
            if _normalize_score_term(score_column_terms.get(score_col)) == term_value
            and _detect_raw_component_bucket(score_col) == component_value
        ]
        ordered_matches = sorted(matching_columns, key=_raw_score_column_sort_key)
        if 1 <= item_number <= len(ordered_matches):
            return ordered_matches[item_number - 1]

    return None


def _build_activity_title_mapping_payload(
    rows: list[dict],
    score_columns: list[str],
    score_column_terms: dict[str, str],
    existing_mappings: dict[str, Any] | None = None,
) -> tuple[dict[str, dict[str, str]], int]:
    keys = list(rows[0].keys()) if rows else []
    score_column_lookup = _build_score_column_lookup(score_columns, score_column_terms)
    payload: dict[str, dict[str, str]] = {}
    unmatched_rows = 0

    for score_col, mapping in (existing_mappings or {}).items():
        if isinstance(mapping, dict) and score_col in score_columns:
            title = str(mapping.get("title") or "").strip()
            if not title:
                continue
            payload[score_col] = {
                "title": title,
                "term": _normalize_score_term(mapping.get("term")) or _normalize_score_term(score_column_terms.get(score_col)) or "",
                "component": _detect_raw_component_bucket(mapping.get("component")) or _detect_raw_component_bucket(score_col) or "",
            }

    for row in rows:
        title_key = _find_column(keys, ["activity title", "activity_title", "title", "activity", "label", "name"])
        title = _normalize_cell(row.get(title_key)).strip() if title_key else ""
        if not title:
            unmatched_rows += 1
            continue

        score_col = _resolve_mapping_target_column(row, keys, score_columns, score_column_terms, score_column_lookup)
        if not score_col:
            unmatched_rows += 1
            continue

        payload[score_col] = {
            "title": title,
            "term": _normalize_score_term(score_column_terms.get(score_col)) or "",
            "component": _detect_raw_component_bucket(score_col) or "",
        }

    return payload, unmatched_rows


def _extract_activity_title_mappings_from_score_keys(
    score_keys: list[str],
    score_column_terms: dict[str, str] | None = None,
) -> dict[str, dict[str, str]]:
    mappings: dict[str, dict[str, str]] = {}

    for score_key in score_keys:
        normalized_key = _normalize_cell(score_key).lower().strip()
        if not normalized_key.startswith(("midterm_", "finals_")):
            continue
        if not any(
            normalized_key.startswith(prefix)
            for prefix in (
                "midterm_class_standing_",
                "midterm_laboratory_",
                "midterm_major_output_",
                "finals_class_standing_",
                "finals_laboratory_",
                "finals_major_output_",
            )
        ):
            continue

        mappings[normalized_key] = {
            "title": _format_activity_title(normalized_key),
            "term": _normalize_score_term((score_column_terms or {}).get(normalized_key)) or ("final" if normalized_key.startswith("finals_") else "midterm"),
            "component": _detect_raw_component_bucket(normalized_key) or "",
        }

    return mappings


def _infer_missing_score_column_terms(score_columns: list[str], existing_terms: dict[str, str]) -> dict[str, str]:
    """Fill missing midterm/final tags for score columns using section-wise ordering heuristics."""
    resolved_terms = {}

    for col in score_columns:
        normalized = _normalize_score_term(existing_terms.get(col))
        if normalized:
            resolved_terms[col] = normalized

    grouped = {
        "class standing": [],
        "laboratory": [],
        "major output": [],
    }

    ordered_cols = sorted(score_columns, key=_raw_score_column_sort_key)
    for col in ordered_cols:
        bucket = _detect_raw_component_bucket(col)
        if bucket in grouped:
            grouped[bucket].append(col)

    for cols in grouped.values():
        if not cols:
            continue

        if len(cols) > 3:
            split_idx = (len(cols) + 1) // 2
            for idx, col in enumerate(cols):
                resolved_terms.setdefault(col, "midterm" if idx < split_idx else "final")
        else:
            for col in cols:
                resolved_terms.setdefault(col, "midterm")

    for col in score_columns:
        if col not in resolved_terms:
            resolved_terms[col] = "final" if _is_final_term_hint_column(col) else "midterm"

    return resolved_terms


def _detect_active_score_source_columns(
    rows: list[dict],
    score_aliases: dict[str, str],
    identity_cols: set[str] | None = None,
) -> set[str]:
    """Detect score source columns that are explicitly populated in score-cap/header rows."""
    identity_cols = set(identity_cols or set())
    if not rows or not score_aliases:
        return set()

    best_columns: set[str] = set()
    best_count = 0

    def _is_likely_index_row(values: list[str]) -> bool:
        """Detect rows like 1..19 that are column index guides, not score-cap rows."""
        if len(values) < 8:
            return False

        ints: list[int] = []
        for raw in values:
            text = _normalize_cell(raw).strip()
            if not re.fullmatch(r"\d+", text):
                return False
            ints.append(int(text))

        if not ints:
            return False
        if min(ints) < 1 or max(ints) > 40:
            return False

        unique_sorted = sorted(set(ints))
        prefix = 0
        for n in unique_sorted:
            if n == prefix + 1:
                prefix += 1
            elif n > prefix + 1:
                break

        repeated_small = sum(1 for n in ints if n <= 5)
        return prefix >= 5 and repeated_small >= 5

    for row in rows[:30]:
        has_identity = any(_normalize_cell(row.get(col, "")).strip() for col in identity_cols)
        if has_identity:
            continue

        candidate_cols: list[str] = []
        candidate_values: list[str] = []
        score_like_count = 0

        for source_col in score_aliases:
            raw = _normalize_cell(row.get(source_col, "")).strip()
            if not raw:
                continue
            candidate_cols.append(source_col)
            candidate_values.append(raw)

            parsed = _to_number_or_text(raw)
            if isinstance(parsed, (int, float)) or _looks_like_raw_fraction_score(raw):
                score_like_count += 1

        if not candidate_cols:
            continue

        # Skip numeric guide rows like 1..19 that are not actual score-cap definitions.
        if _is_likely_index_row(candidate_values):
            continue

        # Prefer rows where most populated cells look like raw score numbers/caps.
        if score_like_count < max(1, len(candidate_cols) // 2):
            continue

        if len(candidate_cols) > best_count:
            best_count = len(candidate_cols)
            best_columns = set(candidate_cols)

    return best_columns


_RAW_SCORE_BUCKET_ORDER = {
    "class standing": 0,
    "laboratory": 1,
    "major output": 2,
}


def _raw_score_column_sort_key(column_name: str) -> tuple[int, int, str]:
    """Sort key: Class Standing 1..n, then Laboratory 1..n, then Major Output 1..n."""
    normalized = _normalize_cell(column_name).lower().strip()

    alias_match = re.fullmatch(r"(class standing|laboratory|major output) item (\d+)", normalized)
    if alias_match:
        bucket = alias_match.group(1)
        item_no = int(alias_match.group(2))
        return (_RAW_SCORE_BUCKET_ORDER.get(bucket, 99), item_no, normalized)

    bucket = _detect_raw_component_bucket(normalized)
    bucket_order = _RAW_SCORE_BUCKET_ORDER.get(bucket, 99)

    if bucket and normalized == bucket:
        return (bucket_order, 1, normalized)

    num_match = re.search(r"(?:item\s*)?(\d+)", normalized)
    item_no = int(num_match.group(1)) if num_match else 999
    return (bucket_order, item_no, normalized)


def _legacy_grade_score_map(enrollment: dict) -> dict:
    """Build a raw component score map from scalar legacy fields when grades_breakdown is missing."""
    score_map = {}
    field_to_column = [
        ("class_standing", "cs (30%)"),
        ("laboratory", "lab (30%)"),
        ("major_output", "mo (40%)"),
        ("final_class_standing", "cs (30%)_2"),
        ("final_laboratory", "lab (30%)_2"),
        ("final_major_output", "mo (40%)_2"),
    ]

    for field, column_name in field_to_column:
        value = enrollment.get(field)
        if value is None or value == "":
            continue
        score_map[column_name] = value

    return score_map


_GRADESHEET_FIELD_MAPPINGS = {
    'class_standing': ['class standing', 'cs (0%)', 'cs (30%)', 'class stand', 'standing', 'cs'],
    'laboratory': ['lab (30%)', 'lab', 'laboratory', 'lab grade'],
    'major_output': ['mo(40%)', 'mo (40%)', 'major output', 'major', 'mo', 'project'],
    'summary': ['summary', 'summ', 'attendance', 'attend'],
    'midterm_grade': ['mtg', 'midterm grade', 'midterm', 'mid', 'mt'],
    'final_grade': ['ftg', 'final grade', 'finals', 'final', 'fin'],
    'overall_grade': ['fg'],
    'final_class_standing': ['cs (30%)_2', 'cs (30%)', 'cs_2', 'cs2'],
    'final_laboratory': ['lab (30%)_2', 'lab_2', 'lab2', 'laboratory_2'],
    'final_major_output': ['mo(40%)', 'mo (40%)', 'mo (40%)_2', 'mo_2', 'mo2', 'major output_2'],
    'midterm_weighted': ['mtg(1/3)'],
    'final_weighted': ['ftg(2/3)'],
    'section_code': ['section', 'section code', 'sec'],
    'subject_code': ['subject code', 'subject_code', 'subjectcode', 'subject', 'course code', 'course_code', 'coursecode', 'code', 'course'],
    'class_time': ['time', 'class time', 'schedule', 'sched'],
}

def _build_gradesheet_update_payload(row, keys, score_aliases, score_alias_terms, active_score_source_cols, student_id=None):
    """Build grade update payloads that replace stale data on re-upload."""
    update_data = {}
    unset_data = {}

    for db_field, keywords in _GRADESHEET_FIELD_MAPPINGS.items():
        col_name = db_field if db_field in row else _find_column(keys, keywords)
        raw_value = _normalize_cell(row.get(col_name, "")) if col_name else ""

        if not raw_value:
            unset_data[db_field] = ""
            continue

        if db_field in {'id_number', 'section_code', 'subject_code', 'class_time'}:
            parsed = _normalize_cell(raw_value)
        else:
            parsed = _to_number_or_text(raw_value)

        if parsed is None or parsed == '':
            unset_data[db_field] = ""
            continue

        update_data[db_field] = parsed

    normalized_student_id = _normalize_student_id(student_id)
    if normalized_student_id:
        update_data['id_number'] = normalized_student_id

    score_map = {}
    score_term_map = {}
    for source_col, score_col in score_aliases.items():
        if active_score_source_cols and source_col not in active_score_source_cols:
            continue
        raw = _normalize_cell(row.get(source_col, ""))
        if not raw:
            continue
        parsed = _to_number_or_text(raw)
        if parsed is not None and _should_include_score_column(score_col, raw, parsed):
            score_map[score_col] = parsed
            score_term_map[score_col] = score_alias_terms.get(source_col, "midterm")

    if score_map:
        update_data['grades_breakdown'] = score_map
        update_data['grades_column_order'] = list(score_map.keys())
        update_data['grades_column_terms'] = score_term_map
    else:
        unset_data['grades_breakdown'] = ""
        unset_data['grades_column_order'] = ""
        unset_data['grades_column_terms'] = ""

    computed_gpa = update_data.get('overall_grade')
    if not isinstance(computed_gpa, (int, float)):
        computed_gpa = update_data.get('final_grade')
    if isinstance(computed_gpa, (int, float)):
        update_data['gpa'] = computed_gpa
    else:
        unset_data['gpa'] = ""

    for field in list(unset_data.keys()):
        if field in update_data:
            unset_data.pop(field, None)

    return update_data, unset_data


def _should_include_score_column(column_key: str, raw_value: str, parsed_value) -> bool:
    """Decide whether a gradesheet column should be treated as a score column."""
    if not column_key:
        return False

    normalized = column_key.lower().strip()
    if normalized.startswith("col_"):
        return False

    if _is_grade_output_column(normalized):
        return False

    # Keep only raw Class Standing/Lab/MO component columns.
    if _is_raw_component_score_column(normalized):
        return True

    # Exclude common non-score metadata columns.
    non_score_exact = {
        "email", "student email",
        "name", "full name", "student name", "name of students", "name of student",
        "id", "id number", "student id", "student no", "student no.", "sid",
        "section", "section code", "sec",
        "subject", "subject code", "course", "course code",
        "class time", "time", "schedule", "sched",
    }
    if normalized in non_score_exact:
        return False

    if any(token in normalized for token in ("email", "section", "subject code", "course code", "class time", "schedule")):
        return False

    value_text = _normalize_cell(raw_value)

    # Always keep explicit raw fraction scores such as 10/20.
    if _looks_like_raw_fraction_score(value_text):
        return True

    # Keep common raw-score assessment columns when they are not aggregate grades.
    raw_assessment_tokens = (
        "quiz", "activity", "assignment", "seatwork", "seat work",
        "recitation", "exercise", "experiment", "project", "output", "item",
        "score", "raw",
    )
    if any(token in normalized for token in raw_assessment_tokens):
        return True

    # Numeric values are likely raw scores unless they look like grade-point outputs.
    if isinstance(parsed_value, (int, float)):
        return True

    return False


def _normalize_subject_code_candidate(value: str) -> str | None:
    candidate = _normalize_cell(value).strip().rstrip(".,;:")
    candidate = re.sub(r"\s+", " ", candidate)
    if not candidate:
        return None
    if not re.search(r"\d", candidate):
        return None
    if re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_./\- ]{1,40}", candidate):
        return candidate
    return None


def _normalize_section_code_candidate(value: str) -> str | None:
    candidate = _normalize_cell(value).strip().rstrip(".,;:")
    candidate = re.sub(r"\s+", " ", candidate)
    if not candidate:
        return None
    if not re.search(r"[A-Za-z0-9]", candidate):
        return None
    if re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_./\- ]{0,30}", candidate):
        return candidate
    return None


def _normalize_subject_name_candidate(value: str) -> str | None:
    candidate = _normalize_cell(value).strip().rstrip(".,;:")
    candidate = re.sub(r"\s+", " ", candidate)
    return candidate or None


def _extract_subject_code_from_text(text: str) -> str | None:
    if not text:
        return None
    match = _SUBJECT_CODE_PATTERN.search(text)
    if match:
        return _normalize_subject_code_candidate(match.group(1))
    return None


def _extract_subject_code_from_values(values: list[str]) -> str | None:
    # Pattern 1: same cell contains "Course Code: ETCH423A"
    for value in values:
        extracted = _extract_subject_code_from_text(value)
        if extracted:
            return extracted

    # Pattern 2: one cell is label and next cell is the code value
    for idx, value in enumerate(values):
        normalized = _normalize_cell(value).lower().strip().rstrip(":")
        if normalized in ("course code", "subject code", "course", "subject"):
            for candidate in values[idx + 1: idx + 4]:
                extracted = _normalize_subject_code_candidate(candidate)
                if extracted:
                    return extracted

    return None


def _extract_labeled_value_from_values(values: list[str], labels: tuple[str, ...], normalizer) -> str | None:
    normalized_labels = {label.lower().strip().rstrip(":") for label in labels}

    for value in values:
        text = _normalize_cell(value).strip()
        if not text or ":" not in text:
            continue
        label_part, candidate_part = text.split(":", 1)
        normalized_label = label_part.lower().strip().rstrip(":")
        if normalized_label in normalized_labels:
            normalized_candidate = normalizer(candidate_part)
            if normalized_candidate:
                return normalized_candidate

    for idx, value in enumerate(values):
        normalized = _normalize_cell(value).lower().strip().rstrip(":")
        if normalized in normalized_labels:
            for candidate in values[idx + 1: idx + 4]:
                normalized_candidate = normalizer(candidate)
                if normalized_candidate:
                    return normalized_candidate

    return None


def _extract_section_code_from_file(file_path: Path) -> str | None:
    return _extract_labeled_metadata_from_file(
        file_path,
        labels=("section", "section code", "sec"),
        normalizer=_normalize_section_code_candidate,
    )


def _extract_subject_name_from_file(file_path: Path) -> str | None:
    return _extract_labeled_metadata_from_file(
        file_path,
        labels=("subject name", "course name", "descriptive title", "course description"),
        normalizer=_normalize_subject_name_candidate,
    )


def _extract_labeled_metadata_from_file(file_path: Path, labels: tuple[str, ...], normalizer) -> str | None:
    ext = file_path.suffix.lower()

    if ext == ".xlsx" and _HAS_OPENPYXL:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        try:
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(min_row=1, max_row=50, values_only=True):
                    values = [_normalize_cell(cell) for cell in row]
                    extracted = _extract_labeled_value_from_values(values, labels, normalizer)
                    if extracted:
                        return extracted
        finally:
            wb.close()

    if ext == ".docx" and _HAS_PYTHON_DOCX:
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            extracted = _extract_labeled_value_from_values([paragraph.text], labels, normalizer)
            if extracted:
                return extracted

        for table in doc.tables:
            for row in table.rows:
                values = [_normalize_cell(cell.text) for cell in row.cells]
                extracted = _extract_labeled_value_from_values(values, labels, normalizer)
                if extracted:
                    return extracted

    if ext == ".csv":
        with open(file_path, newline='', encoding='utf-8') as fh:
            reader = csv.reader(fh)
            for index, row in enumerate(reader):
                if index >= 50:
                    break
                values = [_normalize_cell(cell) for cell in row]
                extracted = _extract_labeled_value_from_values(values, labels, normalizer)
                if extracted:
                    return extracted

    return None


def _normalize_metadata_compare(value: str | None) -> str:
    normalized = _normalize_cell(value).strip().lower()
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized


def _validate_upload_file_matches_class(db, file_path: Path, class_doc: dict, *, upload_type: str = "") -> None:
    expected_section_raw = class_doc.get("section_code")
    if not expected_section_raw and class_doc.get("_id"):
        enrollment = db.enrollments.find_one(
            {"class_id": str(class_doc["_id"]), "section_code": {"$exists": True, "$ne": None}}
        )
        expected_section_raw = (enrollment or {}).get("section_code")

    expected_section = _normalize_metadata_compare(expected_section_raw)
    expected_subject_code = _normalize_metadata_compare(class_doc.get("subject_code"))
    expected_subject_name = _normalize_metadata_compare(class_doc.get("subject_name"))

    extracted_section = _normalize_metadata_compare(_extract_section_code_from_file(file_path))
    extracted_subject_code = _normalize_metadata_compare(_extract_subject_code_from_file(file_path))
    extracted_subject_name = _normalize_metadata_compare(_extract_subject_name_from_file(file_path))

    # Class lists are matched using ordered fallback priority:
    # 1) section code, 2) subject code, 3) subject name.
    # Allow the upload when any available metadata field matches; otherwise fail using
    # the highest-priority comparable field for the error message.
    if upload_type == "classlist":
        comparisons = []
        if extracted_section and expected_section:
            comparisons.append(("section", extracted_section == expected_section))
        if extracted_subject_code and expected_subject_code:
            comparisons.append(("subject_code", extracted_subject_code == expected_subject_code))
        if extracted_subject_name and expected_subject_name:
            comparisons.append(("subject_name", extracted_subject_name == expected_subject_name))

        if any(matches for _, matches in comparisons):
            return

        if extracted_section and expected_section:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"File section code does not match this class. Expected '{expected_section_raw}', "
                    f"but found '{_extract_section_code_from_file(file_path)}'."
                ),
            )

        if extracted_subject_code and expected_subject_code:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"File subject code does not match this class. Expected '{class_doc.get('subject_code')}', "
                    f"but found '{_extract_subject_code_from_file(file_path)}'."
                ),
            )

        if extracted_subject_name and expected_subject_name and extracted_subject_name != expected_subject_name:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"File subject name does not match this class. Expected '{class_doc.get('subject_name')}', "
                    f"but found '{_extract_subject_name_from_file(file_path)}'."
                ),
            )
        return

    if extracted_section and expected_section and extracted_section != expected_section:
        raise HTTPException(
            status_code=400,
            detail=(
                f"File section code does not match this class. Expected '{expected_section_raw}', "
                f"but found '{_extract_section_code_from_file(file_path)}'."
            ),
        )

    if extracted_subject_code and expected_subject_code and extracted_subject_code != expected_subject_code:
        raise HTTPException(
            status_code=400,
            detail=(
                f"File subject code does not match this class. Expected '{class_doc.get('subject_code')}', "
                f"but found '{_extract_subject_code_from_file(file_path)}'."
            ),
        )

    if extracted_subject_name and expected_subject_name and extracted_subject_name != expected_subject_name:
        raise HTTPException(
            status_code=400,
            detail=(
                f"File subject name does not match this class. Expected '{class_doc.get('subject_name')}', "
                f"but found '{_extract_subject_name_from_file(file_path)}'."
            ),
        )


def _extract_subject_code_from_file(file_path: Path) -> str | None:
    ext = file_path.suffix.lower()

    if ext == ".xlsx" and _HAS_OPENPYXL:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        try:
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(min_row=1, max_row=50, values_only=True):
                    values = [_normalize_cell(cell) for cell in row]
                    extracted = _extract_subject_code_from_values(values)
                    if extracted:
                        return extracted
        finally:
            wb.close()

    if ext == ".docx" and _HAS_PYTHON_DOCX:
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            extracted = _extract_subject_code_from_text(paragraph.text)
            if extracted:
                return extracted

        for table in doc.tables:
            for row in table.rows:
                values = [_normalize_cell(cell.text) for cell in row.cells]
                extracted = _extract_subject_code_from_values(values)
                if extracted:
                    return extracted

    if ext == ".csv":
        with open(file_path, newline='', encoding='utf-8') as fh:
            reader = csv.reader(fh)
            for index, row in enumerate(reader):
                if index >= 50:
                    break
                values = [_normalize_cell(cell) for cell in row]
                extracted = _extract_subject_code_from_values(values)
                if extracted:
                    return extracted

    return None


def _parse_file_to_rows(file_path: Path, preferred_type: str | None = None) -> list[dict]:
    """Parse CSV/XLSX/DOCX file into list of row dicts with lowercase keys."""
    ext = os.path.splitext(str(file_path))[1].lower()
    rows = []
    if ext == ".csv":
        with open(file_path, newline='', encoding='utf-8') as fh:
            sample = fh.read(4096)
            fh.seek(0)
            try:
                has_header = csv.Sniffer().has_header(sample)
            except csv.Error:
                has_header = False
            fh.seek(0)
            if has_header:
                reader = csv.DictReader(fh)
                for r in reader:
                    rows.append({k.strip().lower(): str(v).strip() if v else '' for k, v in r.items() if k})
            else:
                reader = csv.reader(fh)
                for row in reader:
                    if row:
                        rows.append({f"col_{i}": str(v).strip() for i, v in enumerate(row)})
    elif ext == ".xlsx":
        if not _HAS_OPENPYXL:
            raise HTTPException(status_code=500, detail="openpyxl required for .xlsx files.")
        wb = openpyxl.load_workbook(file_path, data_only=True)
        try:
            if preferred_type:
                best_rows = []
                best_score = -1
                raw_activity_rows: list[dict] = []

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    raw_rows = list(ws.iter_rows(values_only=True))
                    candidate_rows = _rows_to_dicts(raw_rows)
                    if not candidate_rows:
                        continue

                    candidate_score = _score_rows_for_preferred_type(candidate_rows, preferred_type)
                    if candidate_score > best_score:
                        best_score = candidate_score
                        best_rows = candidate_rows

                    if preferred_type == "gradesheet":
                        term_prefix = _resolve_workbook_term_prefix(sheet_name)
                        if term_prefix:
                            raw_activity_rows = _merge_gradesheet_row_sets(
                                raw_activity_rows,
                                _extract_buksu_raw_activity_rows(ws, term_prefix),
                            )

                if best_rows:
                    rows = _merge_gradesheet_row_sets(best_rows, raw_activity_rows)
                else:
                    ws = wb[wb.sheetnames[0]]
                    raw_rows = list(ws.iter_rows(values_only=True))
                    rows = _merge_gradesheet_row_sets(_rows_to_dicts(raw_rows), raw_activity_rows)
            else:
                ws = wb[wb.sheetnames[0]]
                raw_rows = list(ws.iter_rows(values_only=True))
                rows = _rows_to_dicts(raw_rows)
        finally:
            wb.close()
    elif ext == ".docx":
        if not _HAS_PYTHON_DOCX:
            raise HTTPException(status_code=500, detail="python-docx required for .docx files.")
        doc = Document(file_path)
        if not doc.tables:
            raise HTTPException(status_code=400, detail="No tables found in the Word document.")
        table = doc.tables[0]
        raw_rows = []
        for row in table.rows:
            raw_rows.append([cell.text.strip() for cell in row.cells])

        rows = _rows_to_dicts(raw_rows)
    return rows


def _find_column(keys, keywords):
    """Find first key containing any of the keywords (keywords take priority)."""
    # Check keywords in order of priority, returning first match
    for kw in keywords:
        for key in keys:
            if kw in key:
                return key
    return None


_PRIMARY_STUDENT_ID_KEYWORDS = ['id number', 'student id', 'student_id', 'student no.', 'student number', 'student no', 'school id', 'sid']
_FALLBACK_STUDENT_ID_KEYWORDS = ['id', 'number', 'no.', 'no']


def _find_student_id_column(keys):
    """Find the canonical student-id column, preferring 'ID Number' style headers."""
    col = _find_column(keys, _PRIMARY_STUDENT_ID_KEYWORDS)
    if col:
        return col
    return _find_column(keys, _FALLBACK_STUDENT_ID_KEYWORDS)


def _build_full_name(row, keys):
    """Build a full name from first/middle/last name columns, or a single name column."""
    first_col = _find_column(keys, ['first name', 'first_name', 'firstname', 'fname'])
    middle_col = _find_column(keys, ['middle ini', 'middle name', 'middle_name', 'middlename', 'mname'])
    last_col = _find_column(keys, ['last name', 'last_name', 'lastname', 'lname', 'surname'])
    if first_col or last_col:
        parts = []
        if first_col:
            parts.append(_normalize_cell(row.get(first_col, '')))
        if middle_col:
            parts.append(_normalize_cell(row.get(middle_col, '')))
        if last_col:
            parts.append(_normalize_cell(row.get(last_col, '')))
        full = ' '.join(p for p in parts if p)
        return full or None
    # Fallback: single "name" column or "name of students"
    name_col = _find_column(keys, ['name of students', 'name', 'full name'])
    if name_col:
        return _normalize_cell(row.get(name_col, '')) or None
    return None


def _extract_student_identity(row, keys):
    """Extract student identity from a row with best-effort fallbacks."""
    email_col = _find_column(keys, ['email'])
    # Treat "ID Number" as the canonical file column for student identifiers.
    # Fallback aliases remain for compatibility, but broad matches like "id" or "no."
    # should only be used after the more exact student-id columns fail.
    id_col = _find_student_id_column(keys)
    name_col = _find_column(keys, ['name of students', 'name of student', 'student name', 'student_name', 'full name', 'name'])

    student_email = _normalize_cell(row.get(email_col, '')).lower() if email_col else ''
    student_name = _build_full_name(row, keys)
    if not student_name and name_col:
        student_name = _normalize_cell(row.get(name_col, ''))
    student_id = _normalize_cell(row.get(id_col, '')) if id_col else ''

    return (
        student_email or None,
        student_name or None,
        student_id or None,
    )


def _normalize_student_id(value) -> str:
    return _normalize_cell(value).replace(" ", "").strip()


def _canonicalize_person_name(value: str | None) -> str:
    text = _normalize_cell(value)
    if not text:
        return ""
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = ascii_text.lower()
    ascii_text = re.sub(r"[^a-z0-9,\s]", " ", ascii_text)
    ascii_text = re.sub(r"\s+", " ", ascii_text).strip()
    return ascii_text


def _build_name_match_variants(value: str | None) -> set[str]:
    canonical = _canonicalize_person_name(value)
    if not canonical:
        return set()

    variants = {canonical}
    if "," in canonical:
        last, rest = canonical.split(",", 1)
        reordered = f"{rest.strip()} {last.strip()}".strip()
        reordered = re.sub(r"\s+", " ", reordered)
        if reordered:
            variants.add(reordered)
    return {variant for variant in variants if variant}


def _find_matching_enrollment(db, class_id: str, row, keys):
    """Find an enrollment using student_id, then student_name, then email."""
    student_email, student_name, student_id = _extract_student_identity(row, keys)
    student_id = _normalize_student_id(student_id)

    enrollment = None
    lookup_identifier = None

    if student_id:
        enrollment = db.enrollments.find_one({
            "class_id": class_id,
            "$or": [
                {"student_id": student_id},
                {"id_number": student_id},
                {"student_id": _normalize_cell(student_id)},
                {"id_number": _normalize_cell(student_id)},
            ],
        })
        lookup_identifier = student_id

    if not enrollment and student_name:
        enrollment = db.enrollments.find_one({
            "class_id": class_id,
            "student_name": {"$regex": f"^{re.escape(student_name)}$", "$options": "i"},
        })
        lookup_identifier = student_name

    if not enrollment and student_name:
        target_variants = _build_name_match_variants(student_name)
        for candidate in db.enrollments.find({"class_id": class_id}):
            candidate_email, candidate_name, candidate_id, _candidate_key = _get_enrollment_identity(candidate)
            resolved_candidate_name = _enrich_student_name(db, candidate_email, candidate_id, candidate_name)
            candidate_variants = _build_name_match_variants(resolved_candidate_name or candidate_name)
            if target_variants and candidate_variants and target_variants.intersection(candidate_variants):
                enrollment = candidate
                lookup_identifier = student_name
                break

    if not enrollment and student_email:
        enrollment = db.enrollments.find_one({"class_id": class_id, "student_email": student_email})
        lookup_identifier = student_email

    return enrollment, lookup_identifier, student_email, student_name, student_id


def _row_identifier_label(row, keys) -> str | None:
    student_email, student_name, student_id = _extract_student_identity(row, keys)
    return student_id or student_name or student_email or None


def _get_enrollment_identity(doc):
    """Return normalized identity fields for an enrollment doc."""
    student_email = (doc.get("student_email") or "").strip()
    student_name = (doc.get("student_name") or "").strip()
    student_id = _normalize_student_id(doc.get("student_id") or doc.get("id_number") or "")
    student_key = student_email or student_id or student_name or str(doc.get("_id", ""))
    return student_email, student_name, student_id, student_key


def _enrich_student_name(db, student_email: str, student_id: str, student_name: str) -> str:
    """Best-effort student name resolution using students collection."""
    if student_name:
        return student_name

    try:
        student_doc = None
        if student_email:
            student_doc = db.students.find_one({"email": student_email})
        if not student_doc and student_id:
            student_doc = db.students.find_one({"id_number": student_id})
        if student_doc and student_doc.get("name"):
            return student_doc.get("name")
    except Exception:
        pass

    return student_name


def _is_daily_attendance_format(keys):
    """Detect if file is daily attendance format (checkmarks/signatures instead of percentages)."""
    attendance_keywords = ['date', 'signature', 'face-to-face', 'f2f', 'synchronous', 'asynchronous', 'async']
    has_attendance_markers = any(_find_column(keys, [kw]) for kw in attendance_keywords)
    
    # Check if there are NO month keywords (not monthly format)
    months = ['january', 'february', 'march', 'april', 'may', 'june',
              'july', 'august', 'september', 'october', 'november', 'december']
    has_months = any(_find_column(keys, [month]) for month in months)
    
    # Daily format has date/signature markers but NOT month names
    return has_attendance_markers and not has_months


def _parse_daily_attendance(row, keys, id_col, name_col):
    """Parse daily attendance from checkmarks.
    
    Returns: (present_days, absent_days, attendance_percentage, marked_columns)
    """
    marked_columns = []
    present_days = 0
    total_attendance_cols = 0
    
    # Determine which columns are attendance (skip ID and Name)
    name_keywords = ['name of students', 'name of student', 'name', 'student name', 'student_name']
    id_keywords = ['no.', 'no', 'number', 'id', 'student id', 'id number']
    
    for col_name in keys:
        # Skip ID and Name columns
        is_name = any(kw in col_name.lower() for kw in name_keywords)
        is_id = any(kw in col_name.lower() for kw in id_keywords)
        if is_name or is_id or col_name == id_col or col_name == name_col:
            continue
        
        # Check if column has any attendance data (not empty)
        cell_value = row.get(col_name, '').strip()
        if cell_value or col_name in [k for k in keys if k not in ['name of students', id_col, name_col]]:
            total_attendance_cols += 1
            if '✓' in cell_value or cell_value.lower() in ['present', 'p', 'yes']:
                present_days += 1
                marked_columns.append(col_name)
    
    # Calculate attendance percentage
    absent_days = total_attendance_cols - present_days if total_attendance_cols > 0 else 0
    attendance_pct = (present_days / total_attendance_cols * 100) if total_attendance_cols > 0 else 0
    
    return present_days, absent_days, round(attendance_pct, 2), marked_columns


# --- Preview classlist (extract students without saving) ---
@router.post("/preview-classlist", status_code=200)
async def preview_classlist(
    actor: dict = Depends(get_current_actor),
    class_id: str | None = Form(None),
    file: UploadFile = File(..., description="CSV or XLSX file"),
):
    """Parse a classlist file and return extracted students (name and ID) without saving."""
    try:
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in [".csv", ".xlsx"]:
            raise HTTPException(
                status_code=400,
                detail="Only CSV and XLSX files are supported for preview."
            )

        # Save file temporarily
        buffer = io.BytesIO()
        content = await file.read()
        buffer.write(content)
        buffer.seek(0)
        
        from tempfile import NamedTemporaryFile
        with NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            if class_id:
                if not ObjectId.is_valid(class_id):
                    raise HTTPException(status_code=404, detail="Class not found")
                db = get_db()
                class_doc = _get_class_for_actor(db, class_id, actor)
                _validate_upload_file_matches_class(db, tmp_path, class_doc, upload_type="classlist")

            # Parse the file
            rows = _parse_file_to_rows(tmp_path, preferred_type="classlist")
            if not rows:
                raise HTTPException(status_code=400, detail="No data found in file.")

            keys = list(rows[0].keys()) if rows else []
            
            # Extract students
            students = []
            for row in rows:
                student_email, student_name, student_id = _extract_student_identity(row, keys)
                student_id = _normalize_student_id(student_id)
                
                # Only include rows with at least a name or ID
                if student_name or student_id:
                    students.append({
                        "name": student_name or "Unknown",
                        "id": student_id or "N/A",
                        "email": student_email or None
                    })

            return {
                "students": students,
                "count": len(students),
                "parsed_columns": keys
            }

        finally:
            # Clean up temp file with retry logic
            import time
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    if tmp_path.exists():
                        tmp_path.unlink()
                    break
                except (FileNotFoundError, PermissionError):
                    if attempt < max_retries - 1:
                        time.sleep(0.1)  # Small delay before retry
                    elif tmp_path.exists():
                        # If still locked after retries, try to open it to clear the lock
                        try:
                            open(str(tmp_path), 'rb').close()
                            tmp_path.unlink()
                        except Exception:
                            pass  # File will be cleaned up by OS temp cleanup

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")


# --- Upload gradesheet/attendance/classlist for a class ---
@router.post("/{class_id}/upload", status_code=201)
async def upload_class_files(
    class_id: str,
    actor: dict = Depends(get_current_actor),
    files: List[UploadFile] = File(..., description="CSV, XLSX, or DOCX files"),
    type: str = Form(...)
):
    """Upload classlist, gradesheet, or attendance for a class. Parses and saves all data to DB."""
    if not ObjectId.is_valid(class_id):
        raise HTTPException(status_code=404, detail="Class not found")
    if type not in ("gradesheet", "attendance", "classlist"):
        raise HTTPException(status_code=400, detail="Invalid upload type.")
    saved_files = []
    # classlist tracking
    add_summary = {"added": 0, "skipped": 0, "invalid": 0}
    added_students = []
    skipped_students = []
    invalid_students = []
    # gradesheet/attendance tracking
    updated_count = 0
    not_enrolled = []
    missing_identifiers = 0
    auto_referred_students = []
    auto_activity_title_mappings: dict[str, dict[str, str]] = {}

    try:
        db = get_db()
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")

    class_doc = db.classes.find_one({"_id": ObjectId(class_id)})
    if not class_doc:
        raise HTTPException(status_code=404, detail="Class not found")

    for upload in files:
        ext = os.path.splitext(upload.filename)[1].lower()
        if ext not in [".csv", ".xlsx", ".docx"]:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Please upload CSV, XLSX, or DOCX files only.")
        unique_name = f"{class_id}_{type}_{uuid.uuid4().hex}{ext}"
        dest = UPLOAD_DIR / unique_name
        with dest.open("wb") as buffer:
            shutil.copyfileobj(upload.file, buffer)
        saved_files.append(str(dest.name))

        if type in ("gradesheet", "attendance", "classlist"):
            _validate_upload_file_matches_class(db, dest, class_doc, upload_type=type)

        try:
            rows = _parse_file_to_rows(dest, preferred_type=type)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse {upload.filename}: {exc}")

        if not rows:
            raise HTTPException(status_code=400, detail=f"{upload.filename} did not contain any readable rows.")

        keys = list(rows[0].keys())
        if type == "gradesheet":
            detected_type = _detect_file_type(keys)
            if detected_type == "attendance":
                raise HTTPException(
                    status_code=400,
                    detail="Attendance files should be uploaded from the Attendance page, not the Grades page.",
                )
        email_col = _find_column(keys, ['email'])
        # Detect name columns for excluding from extra data
        _name_cols = set()
        for kw in ['first name', 'first_name', 'firstname', 'fname',
                    'middle ini', 'middle name', 'middle_name', 'middlename', 'mname',
                    'last name', 'last_name', 'lastname', 'lname', 'surname', 'name']:
            col = _find_column(keys, [kw])
            if col:
                _name_cols.add(col)

        if type == "classlist":
            for row in rows:
                student_email, student_name, student_id = _extract_student_identity(row, keys)
                student_id = _normalize_student_id(student_id)
                
                section_col = _find_column(keys, ['section', 'section code', 'sec', 'section_code'])
                section_code = _normalize_cell(row.get(section_col, '')) if section_col else None
                
                # Skip rows with no student identity
                if not student_name and not student_id:
                    add_summary['invalid'] += 1
                    invalid_students.append({"email": None, "name": None})
                    continue
                
                # Build student profile data if email is available
                if student_email:
                    student_data = {"email": student_email}
                    if student_name:
                        student_data["name"] = student_name
                    if student_id:
                        student_data["id_number"] = student_id
                    for k, v in row.items():
                        if k == email_col or k in _name_cols or not v:
                            continue
                        student_data[k] = v
                    # Upsert into students collection if email available
                    db.students.update_one({"email": student_email}, {"$set": student_data}, upsert=True)
                
                # Check if enrollment already exists (by student_id or student_name)
                existing = None
                if student_id:
                    existing = db.enrollments.find_one({
                        "class_id": class_id,
                        "$or": [{"student_id": student_id}, {"id_number": student_id}]
                    })
                
                if not existing and student_name:
                    existing = db.enrollments.find_one({
                        "class_id": class_id,
                        "student_name": {"$regex": f"^{re.escape(student_name)}$", "$options": "i"}
                    })
                
                if not existing and student_email:
                    existing = db.enrollments.find_one({"class_id": class_id, "student_email": student_email})
                
                if existing:
                    existing_updates = {}
                    if student_name and _normalize_cell(existing.get("student_name")) != student_name:
                        existing_updates["student_name"] = student_name
                    if student_email and _normalize_cell(existing.get("student_email")).lower() != student_email:
                        existing_updates["student_email"] = student_email
                    if student_id:
                        if _normalize_student_id(existing.get("student_id")) != student_id:
                            existing_updates["student_id"] = student_id
                        if _normalize_student_id(existing.get("id_number")) != student_id:
                            existing_updates["id_number"] = student_id
                    if section_code and _normalize_cell(existing.get("section_code")) != section_code:
                        existing_updates["section_code"] = section_code
                    if existing_updates:
                        db.enrollments.update_one({"_id": existing["_id"]}, {"$set": existing_updates})
                    add_summary['skipped'] += 1
                    skipped_students.append({"email": student_email, "name": student_name})
                else:
                    # Create enrollment
                    enrollment_data = {"class_id": class_id}
                    if student_name:
                        enrollment_data["student_name"] = student_name
                    if student_id:
                        enrollment_data["student_id"] = student_id
                        enrollment_data["id_number"] = student_id
                    if student_email:
                        enrollment_data["student_email"] = student_email
                    if section_code:
                        enrollment_data["section_code"] = section_code
                    
                    db.enrollments.insert_one(enrollment_data)
                    add_summary['added'] += 1
                    added_students.append({"email": student_email or None, "name": student_name})


        elif type == "gradesheet":
            name_col = _find_column(keys, ['name of students', 'name', 'student name', 'student_name', 'full name', 'name of student'])
            id_col = _find_student_id_column(keys)
            identity_cols = {k for k in [email_col, name_col, id_col] if k}
            score_aliases = _build_raw_score_column_aliases(
                keys,
                identity_cols,
                alias_all_section_columns=True,
            )
            score_alias_terms = _build_raw_score_column_terms(
                keys,
                score_aliases,
                identity_cols,
            )
            auto_activity_title_mappings.update(
                _extract_activity_title_mappings_from_score_keys(
                    list(score_aliases.values()),
                    score_alias_terms,
                )
            )
            active_score_source_cols = _detect_active_score_source_columns(
                rows,
                score_aliases,
                identity_cols,
            )

            for row in rows:
                enrollment, lookup_identifier, _email, _name, matched_student_id = _find_matching_enrollment(db, class_id, row, keys)
                if not enrollment:
                    if not _row_identifier_label(row, keys):
                        missing_identifiers += 1
                    if lookup_identifier:
                        not_enrolled.append(lookup_identifier)
                    continue

                update_data, unset_data = _build_gradesheet_update_payload(
                    row,
                    keys,
                    score_aliases,
                    score_alias_terms,
                    active_score_source_cols,
                    matched_student_id,
                )

                if update_data or unset_data:
                    update_ops = {}
                    if update_data:
                        update_ops["$set"] = update_data
                    if unset_data:
                        update_ops["$unset"] = unset_data
                    db.enrollments.update_one(
                        {"_id": enrollment["_id"]},
                        update_ops
                    )
                    refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                    if refreshed:
                        was_referred = bool(enrollment.get("flagged_for_mentoring"))
                        _apply_automatic_referral(db, class_doc, refreshed)
                        refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                        if refreshed and not was_referred and refreshed.get("flagged_for_mentoring"):
                            auto_referred_students.append({
                                "student_name": refreshed.get("student_name"),
                                "student_id": refreshed.get("student_id") or refreshed.get("id_number"),
                                "student_email": refreshed.get("student_email"),
                                "referral_reasons": _build_auto_referral_reasons(refreshed),
                            })
                    updated_count += 1

        elif type == "attendance":
            months = ['january', 'february', 'march', 'april', 'may', 'june', 
                     'july', 'august', 'september', 'october', 'november', 'december']
            
            # Detect if this is daily attendance format (checkmarks) or monthly format
            is_daily_format = _is_daily_attendance_format(keys)
            
            for row in rows:
                name_col = _find_column(keys, ['name of students', 'name', 'student name', 'student_name', 'full name', 'name of student'])
                id_col = _find_student_id_column(keys)
                enrollment, lookup_identifier, _email, _name, matched_student_id = _find_matching_enrollment(db, class_id, row, keys)
                if not enrollment:
                    if not _row_identifier_label(row, keys):
                        missing_identifiers += 1
                    if lookup_identifier:
                        not_enrolled.append(lookup_identifier)
                    continue
                
                update_data = {}
                
                # Handle daily attendance format (checkmarks)
                if is_daily_format:
                    present_days, absent_days, attendance_pct, marked_cols = _parse_daily_attendance(row, keys, id_col, name_col)
                    if present_days > 0 or absent_days > 0:
                        update_data['attendance_present_days'] = present_days
                        update_data['attendance_absent_days'] = absent_days
                        update_data['attendance_overall'] = attendance_pct
                        update_data['attendance'] = attendance_pct
                        update_data['attendance_marked_columns'] = marked_cols
                
                # Handle monthly attendance format (percentages)
                else:
                    monthly_values = []
                    for month in months:
                        col_name = _find_column(keys, [month])
                        raw_month_value = _normalize_cell(row.get(col_name, '')) if col_name else ''
                        if col_name and raw_month_value:
                            val = raw_month_value
                            try:
                                attendance_pct = float(val.replace('%', ''))
                                update_data[f'attendance_{month}'] = attendance_pct
                                monthly_values.append(attendance_pct)
                            except (ValueError, TypeError):
                                pass
                    
                    if monthly_values:
                        update_data['attendance_overall'] = round(sum(monthly_values) / len(monthly_values), 2)
                        update_data['attendance'] = update_data['attendance_overall']
                
                # Extract attendance-related fields
                normalized_student_id = _normalize_student_id(matched_student_id)
                if normalized_student_id:
                    update_data['id_number'] = normalized_student_id
                
                section_col = _find_column(keys, ['section', 'section code', 'sec'])
                if section_col and _normalize_cell(row.get(section_col, '')):
                    update_data['section_code'] = _normalize_cell(row.get(section_col, ''))
                
                subject_col = _find_column(keys, ['subject code', 'subject_code', 'subjectcode', 'subject', 'course code', 'course_code', 'coursecode', 'code', 'course'])
                if subject_col and _normalize_cell(row.get(subject_col, '')):
                    update_data['subject_code'] = _normalize_cell(row.get(subject_col, ''))
                
                if update_data:
                    db.enrollments.update_one(
                        {"_id": enrollment["_id"]},
                        {"$set": update_data}
                    )
                    refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                    if refreshed:
                        was_referred = bool(enrollment.get("flagged_for_mentoring"))
                        _apply_automatic_referral(db, class_doc, refreshed)
                        refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                        if refreshed and not was_referred and refreshed.get("flagged_for_mentoring"):
                            auto_referred_students.append({
                                "student_name": refreshed.get("student_name"),
                                "student_id": refreshed.get("student_id") or refreshed.get("id_number"),
                                "student_email": refreshed.get("student_email"),
                                "referral_reasons": _build_auto_referral_reasons(refreshed),
                            })
                    updated_count += 1

    if type == "gradesheet" and auto_activity_title_mappings:
        db.classes.update_one(
            {"_id": ObjectId(class_id)},
            {
                "$set": {
                    "activity_title_mappings": auto_activity_title_mappings,
                    "activity_title_mapping_updated_at": datetime.now(timezone.utc),
                }
            },
        )
        for enrollment in db.enrollments.find({"class_id": class_id}):
            enriched_enrollment = {
                **enrollment,
                "activity_title_mappings": auto_activity_title_mappings,
            }
            topic_payload = _extract_topic_difficulty(
                enriched_enrollment,
                build_model_feature_dict(enriched_enrollment),
            )
            db.enrollments.update_one(
                {"_id": enrollment["_id"]},
                {
                    "$set": {
                        "activity_title_mappings": auto_activity_title_mappings,
                        "midterm_topic_difficulties": topic_payload.get("midterm_topic_difficulties") or [],
                        "hardest_midterm_topics": topic_payload.get("hardest_midterm_topics") or [],
                    }
                },
            )

    result = {
        "message": f"{type.capitalize()} file(s) uploaded and saved successfully.",
        "files": saved_files,
        "auto_referred_students": auto_referred_students,
    }
    create_activity_log(
        db,
        actor_id=actor["id"],
        actor_name=actor.get("name", "User"),
        role=actor["role"],
        action=f"upload_{type}",
        description=f"Uploaded {type} file(s) for {(class_doc.get('section_code') or '').strip()} {(class_doc.get('subject_code') or '').strip()} {(class_doc.get('subject_name') or '').strip()}".strip(),
        target_type="class",
        target_id=class_id,
        metadata={
            "files": saved_files,
            "updated": updated_count,
            "added": add_summary.get("added", 0),
            "skipped": add_summary.get("skipped", 0),
            "invalid": add_summary.get("invalid", 0),
            "auto_referred_count": len(auto_referred_students),
        },
    )
    if type == "classlist":
        result.update(add_summary)
        result["added_students"] = added_students
        result["skipped_students"] = skipped_students
        result["invalid_students"] = invalid_students
    else:
        result["updated"] = updated_count
        if not_enrolled:
            result["not_enrolled"] = sorted(set(not_enrolled))[:25]
        if missing_identifiers:
            result["missing_identifiers"] = missing_identifiers
    return JSONResponse(result)


def _detect_file_type(keys):
    """Detect if keys represent classlist, gradesheet, or attendance file."""
    # Check for grade-related columns
    grade_keywords = [
        'midterm', 'prelim', 'final', 'lab', 'laboratory', 'standing',
        'major output', 'summary', 'grade', 'gpa', 'quiz', 'exam',
        'score', 'points', 'rating', 'remarks'
    ]
    has_grades = any(_find_column(keys, [kw]) for kw in grade_keywords)

    # Check for attendance markers (month/date/signature style)
    attendance_keywords = [
        'january', 'february', 'march', 'april', 'may', 'june',
        'july', 'august', 'september', 'october', 'november', 'december',
        'attendance', 'date', 'signature', 'present', 'absent',
        'face-to-face', 'f2f', 'synchronous', 'asynchronous'
    ]
    has_attendance = any(_find_column(keys, [kw]) for kw in attendance_keywords)
    
    # Check for name/id columns (classlist)
    name_keywords = ['name of students', 'name', 'first name', 'last name', 'student name', 'student_name', 'full name']
    id_keywords = ['no', 'number', 'no.', 'id', 'student id', 'id number', 'school id']
    has_name_or_id = any(_find_column(keys, [kw]) for kw in name_keywords + id_keywords)
    
    # Determine file type
    if has_grades and not has_attendance:
        return "gradesheet"
    elif has_attendance:
        return "attendance"
    elif has_name_or_id:
        return "classlist"
    else:
        return "unknown"


@router.post("/upload-classlist", status_code=201)
async def upload_and_create_classlist(
    actor: dict = Depends(get_current_actor),
    files: List[UploadFile] = File(..., description="CSV, XLSX, or DOCX files (Classlist, Gradesheet, or Attendance)"),
    instructor_id: str = Form(...),
    subject_code: str = Form(None, description="Optional: Subject/Course code (required if not in file)"),
    course_code: str = Form(None, description="Optional alias for subject_code"),
    subject_name: str = Form(None, description="Optional: Subject/Course name")
):
    """Upload classlist, gradesheet, and/or attendance files. Auto-detects file types and processes accordingly."""
    try:
        db = get_db()
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")

    try:
        instructor_id = str(instructor_id).strip()
        _ensure_instructor_scope(actor, instructor_id)
        if not instructor_id:
            raise HTTPException(status_code=400, detail="Instructor ID is required.")
        
        # Parse and categorize files
        file_data = {}  # { "classlist": [rows], "gradesheet": [rows], "attendance": [rows] }
        saved_files = []
        
        detected_subject_code = None
        detected_section_code = None
        detected_subject_name = None
        
        for upload in files:
            ext = os.path.splitext(upload.filename)[1].lower()
            if ext not in [".csv", ".xlsx", ".docx"]:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
            
            unique_name = f"{instructor_id}_data_{uuid.uuid4().hex}{ext}"
            dest = UPLOAD_DIR / unique_name
            with dest.open("wb") as buffer:
                shutil.copyfileobj(upload.file, buffer)
            saved_files.append(unique_name)

            try:
                rows = _parse_file_to_rows(dest)
            except Exception:
                raise HTTPException(status_code=400, detail=f"Failed to parse {upload.filename}.")

            if not rows:
                continue

            keys = list(rows[0].keys())
            file_type = _detect_file_type(keys)
            
            if file_type is None or file_type == "unknown":
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Could not determine file type for {upload.filename}. "
                        "Ensure it has recognizable columns for classlist (name/id), "
                        "gradesheet (grade/score), or attendance (month/date/signature)."
                    ),
                )
            
            if file_type not in file_data:
                file_data[file_type] = []
            file_data[file_type].extend(rows)
            file_data[file_type + "_keys"] = keys
            
            # Detect class info from any file type that has these columns
            section_col = _find_column(keys, ['section', 'section code', 'sec', 'section_code'])
            subject_col = _find_column(keys, ['subject', 'subject code', 'subject_code', 'subjectcode', 'course code', 'course_code', 'coursecode', 'code', 'course'])
            subject_name_col = _find_column(keys, ['subject name', 'course name', 'course_name', 'subject_name', 'coursename'])
            
            if rows:
                if section_col and not detected_section_code:
                    detected_section_code = rows[0].get(section_col, '').strip()
                if subject_col and not detected_subject_code:
                    detected_subject_code = rows[0].get(subject_col, '').strip()
                if subject_name_col and not detected_subject_name:
                    detected_subject_name = rows[0].get(subject_name_col, '').strip()

            # Fallback: extract "Course Code: XXXX" metadata from the file body.
            if not detected_subject_code:
                extracted_code = _extract_subject_code_from_file(dest)
                if extracted_code:
                    detected_subject_code = extracted_code

        # Use detected subject code, fallback to form parameter
        if not detected_subject_code:
            detected_subject_code = (subject_code or course_code or "").strip() or None
        if not detected_subject_name and subject_name:
            detected_subject_name = subject_name
        
        if not detected_subject_code:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not detect subject code. The system treats course code as subject code. "
                    "Ensure the file has 'course code'/'subject code' or provide subject_code/course_code parameter."
                ),
            )

        # Build class roster source(s): classlist OR gradesheet OR attendance,
        # as long as at least one has student identity columns.
        seed_sources = []
        for source_type in ("classlist", "gradesheet", "attendance"):
            source_rows = file_data.get(source_type) or []
            source_keys = file_data.get(f"{source_type}_keys") or []
            if not source_rows or not source_keys:
                continue

            has_name = _find_column(source_keys, ['name of students', 'name of student', 'student name', 'student_name', 'full name', 'name', 'first name', 'last name']) is not None
            has_id = _find_student_id_column(source_keys) is not None
            has_email = _find_column(source_keys, ['email']) is not None

            if has_name or has_id or has_email:
                seed_sources.append((source_type, source_rows, source_keys))

        if not seed_sources:
            raise HTTPException(
                status_code=400,
                detail="At least one uploaded file must include student name, student id, or email columns to build the class list.",
            )

        # Create roster summary
        classlist_summary = {"added": 0, "skipped": 0, "invalid": 0}
        classlist_added = []
        classlist_skipped = []
        classlist_invalid = []

        # Create or find class
        class_doc = db.classes.find_one({
            "instructor_id": instructor_id,
            "subject_code": detected_subject_code,
            "status": "active"
        })
        
        if not class_doc:
            class_data = {
                "instructor_id": instructor_id,
                "subject_code": detected_subject_code,
                "subject_name": detected_subject_name or detected_subject_code,
                "status": "active",
            }
            result = db.classes.insert_one(class_data)
            class_id = str(result.inserted_id)
        else:
            class_id = str(class_doc["_id"])

        # Build students + enrollments from any seed source.
        seen_seed_keys = set()
        for source_type, source_rows, source_keys in seed_sources:
            source_email_col = _find_column(source_keys, ['email'])
            source_id_col = _find_student_id_column(source_keys)
            source_section_col = _find_column(source_keys, ['section', 'section code', 'sec', 'section_code'])

            source_name_cols = set()
            for kw in ['first name', 'first_name', 'firstname', 'fname',
                        'middle ini', 'middle name', 'middle_name', 'middlename', 'mname',
                        'last name', 'last_name', 'lastname', 'lname', 'surname',
                        'name of students', 'name of student', 'student name', 'student_name', 'full name', 'name']:
                col = _find_column(source_keys, [kw])
                if col:
                    source_name_cols.add(col)

            for row in source_rows:
                student_email, student_name, student_id = _extract_student_identity(row, source_keys)
                student_id = _normalize_student_id(student_id)
                section_code = row.get(source_section_col, '').strip() if source_section_col else None

                if student_id:
                    dedupe_key = f"id:{student_id.lower()}"
                elif student_name:
                    dedupe_key = f"name:{student_name.lower()}"
                elif student_email:
                    dedupe_key = f"email:{student_email.lower()}"
                else:
                    classlist_summary['invalid'] += 1
                    classlist_invalid.append({"source": source_type, "name": None, "id": None})
                    continue

                if dedupe_key in seen_seed_keys:
                    continue
                seen_seed_keys.add(dedupe_key)

                # Upsert student profile
                student_data = {}
                if student_name:
                    student_data["name"] = student_name
                if student_id:
                    student_data["id_number"] = student_id
                if student_email:
                    student_data["email"] = student_email

                for k, v in row.items():
                    if not v:
                        continue
                    if k in source_name_cols or k == source_email_col or k == source_id_col:
                        continue
                    student_data[k] = v

                if student_id:
                    student_filter = {"id_number": student_id}
                elif student_name:
                    student_filter = {"name": student_name}
                else:
                    student_filter = {"email": student_email}
                db.students.update_one(student_filter, {"$set": student_data}, upsert=True)

                # Upsert enrollment in this class
                existing = None
                if student_id:
                    existing = db.enrollments.find_one({
                        "class_id": class_id,
                        "$or": [{"student_id": student_id}, {"id_number": student_id}],
                    })
                if not existing and student_name:
                    existing = db.enrollments.find_one({
                        "class_id": class_id,
                        "student_name": {"$regex": f"^{re.escape(student_name)}$", "$options": "i"},
                    })
                if not existing and student_email:
                    existing = db.enrollments.find_one({"class_id": class_id, "student_email": student_email})

                if existing:
                    existing_updates = {}
                    if student_name and _normalize_cell(existing.get("student_name")) != student_name:
                        existing_updates["student_name"] = student_name
                    if student_email and _normalize_cell(existing.get("student_email")).lower() != (student_email or ""):
                        existing_updates["student_email"] = student_email
                    if student_id:
                        if _normalize_student_id(existing.get("student_id")) != student_id:
                            existing_updates["student_id"] = student_id
                        if _normalize_student_id(existing.get("id_number")) != student_id:
                            existing_updates["id_number"] = student_id
                    if section_code and _normalize_cell(existing.get("section_code")) != section_code:
                        existing_updates["section_code"] = section_code
                    if existing_updates:
                        db.enrollments.update_one({"_id": existing["_id"]}, {"$set": existing_updates})
                    classlist_summary['skipped'] += 1
                    classlist_skipped.append({"source": source_type, "name": student_name, "id": student_id})
                else:
                    enrollment_data = {"class_id": class_id}
                    if student_name:
                        enrollment_data["student_name"] = student_name
                    if student_id:
                        enrollment_data["student_id"] = student_id
                        enrollment_data["id_number"] = student_id
                    if student_email:
                        enrollment_data["student_email"] = student_email
                    if section_code:
                        enrollment_data["section_code"] = section_code
                    db.enrollments.insert_one(enrollment_data)
                    classlist_summary['added'] += 1
                    classlist_added.append({"source": source_type, "name": student_name, "id": student_id})

        # Process gradesheet if present
        gradesheet_summary = {}
        if "gradesheet" in file_data and file_data["gradesheet"]:
            gradesheet_rows = file_data["gradesheet"]
            gradesheet_keys = file_data["gradesheet_keys"]
            updated_count_gs = 0
            not_enrolled_gs = []
            gs_email_col = _find_column(gradesheet_keys, ['email'])
            gs_name_col = _find_column(gradesheet_keys, ['name of students', 'name of student', 'student name', 'student_name', 'full name', 'name'])
            gs_id_col = _find_student_id_column(gradesheet_keys)
            gs_identity_cols = {k for k in [gs_email_col, gs_name_col, gs_id_col] if k}
            gs_score_aliases = _build_raw_score_column_aliases(
                gradesheet_keys,
                gs_identity_cols,
                alias_all_section_columns=True,
            )
            gs_score_alias_terms = _build_raw_score_column_terms(
                gradesheet_keys,
                gs_score_aliases,
                gs_identity_cols,
            )
            gs_active_score_source_cols = _detect_active_score_source_columns(
                gradesheet_rows,
                gs_score_aliases,
                gs_identity_cols,
            )

            for row in gradesheet_rows:
                enrollment, lookup_identifier, _email, _name, _id = _find_matching_enrollment(db, class_id, row, gradesheet_keys)
                if not enrollment:
                    if lookup_identifier:
                        not_enrolled_gs.append(lookup_identifier)
                    continue

                update_data, unset_data = _build_gradesheet_update_payload(
                    row,
                    gradesheet_keys,
                    gs_score_aliases,
                    gs_score_alias_terms,
                    gs_active_score_source_cols,
                )

                if update_data or unset_data:
                    update_ops = {}
                    if update_data:
                        update_ops["$set"] = update_data
                    if unset_data:
                        update_ops["$unset"] = unset_data
                    db.enrollments.update_one(
                        {"_id": enrollment["_id"]},
                        update_ops
                    )
                    refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                    if refreshed:
                        _apply_automatic_referral(db, class_doc, refreshed)
                    updated_count_gs += 1

            gradesheet_summary = {"updated": updated_count_gs, "not_enrolled": len(not_enrolled_gs)}

        # Process attendance if present
        attendance_summary = {}
        if "attendance" in file_data and file_data["attendance"]:
            attendance_rows = file_data["attendance"]
            attendance_keys = file_data["attendance_keys"]
            updated_count_att = 0
            not_enrolled_att = []

            month_keywords = ['january', 'february', 'march', 'april', 'may', 'june',
                              'july', 'august', 'september', 'october', 'november', 'december']

            for row in attendance_rows:
                enrollment, lookup_identifier, _email, _name, _id = _find_matching_enrollment(db, class_id, row, attendance_keys)
                if not enrollment:
                    if lookup_identifier:
                        not_enrolled_att.append(lookup_identifier)
                    continue

                update_data = {}
                overall_total = 0
                overall_count = 0

                for month in month_keywords:
                    month_col = _find_column(attendance_keys, [month])
                    if month_col and row.get(month_col, '').strip():
                        try:
                            val = float(row.get(month_col, '').strip())
                            update_data[month] = val
                            overall_total += val
                            overall_count += 1
                        except (ValueError, TypeError):
                            pass

                if overall_count > 0:
                    update_data['attendance'] = overall_total / overall_count

                if update_data:
                    db.enrollments.update_one(
                        {"_id": enrollment["_id"]},
                        {"$set": update_data}
                    )
                    refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
                    if refreshed:
                        _apply_automatic_referral(db, class_doc, refreshed)
                    updated_count_att += 1

            attendance_summary = {"updated": updated_count_att, "not_enrolled": len(not_enrolled_att)}

        return {
            "class_id": class_id,
            "subject_code": detected_subject_code,
            "subject_name": detected_subject_name or detected_subject_code,
            "section_code": detected_section_code,
            "classlist_summary": classlist_summary,
            "classlist_added": classlist_added[:10],
            "classlist_skipped": classlist_skipped[:10],
            "classlist_invalid": classlist_invalid[:10],
            "gradesheet_summary": gradesheet_summary if gradesheet_summary else None,
            "attendance_summary": attendance_summary if attendance_summary else None,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


def _doc_to_class_response(doc, student_count: int = 0, at_risk_count: int = 0, section_code: str = None) -> dict:
    return {
        "id": str(doc["_id"]),
        "subject_code": doc["subject_code"],
        "subject_name": doc["subject_name"],
        "instructor_id": doc["instructor_id"],
        "status": doc.get("status", "active"),
        "section_code": section_code,
        "student_count": student_count,
        "at_risk_count": at_risk_count,
    }


@router.get("/risk-alerts")
def list_instructor_risk_alerts(instructor_id: str, actor: dict = Depends(get_current_actor)):
    """List all referred students across the instructor's classes (for alerts page)."""
    try:
        db = get_db()
        classes_cursor = db.classes.find({"instructor_id": instructor_id}).sort("subject_code", 1)
        alerts = []
        for c in classes_cursor:
            class_id = str(c["_id"])
            subject_code = c.get("subject_code", "")
            subject_name = c.get("subject_name", "")
            cursor = db.enrollments.find(
                {"class_id": class_id, "flagged_for_mentoring": True}
            ).sort("student_email", 1)
            for doc in cursor:
                student_email, student_name, student_id, student_key = _get_enrollment_identity(doc)
                resolved_name = _enrich_student_name(db, student_email, student_id, student_name)
                alerts.append({
                    "student_email": student_email or student_key,
                    "student_name": resolved_name or None,
                    "student_id": student_id or None,
                    "prediction_label": "External Factor" if doc.get("risk_source") == "external_factors" else ("Academic Problem" if doc.get("risk_source") else None),
                    "class_id": class_id,
                    "subject_code": subject_code,
                    "subject_name": subject_name,
                })
        return alerts
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/instructor-students")
def list_instructor_students(instructor_id: str, actor: dict = Depends(get_current_actor)):
    """List all students (enrollments) across the instructor's classes for the Student List page."""
    try:
        db = get_db()
        classes_cursor = db.classes.find({"instructor_id": instructor_id}).sort("subject_code", 1)
        rows = []
        for c in classes_cursor:
            class_id = str(c["_id"])
            subject_code = c.get("subject_code", "")
            subject_name = c.get("subject_name", "")
            # Sort by _id instead of student_email which may not exist
            cursor = db.enrollments.find({"class_id": class_id}).sort("_id", 1)
            for doc in cursor:
                student_email, student_name, student_id, student_key = _get_enrollment_identity(doc)
                resolved_name = _enrich_student_name(db, student_email, student_id, student_name)
                row = {
                    "student_email": student_email or student_key,
                    "student_name": resolved_name or student_name or None,
                    "student_id": student_id or None,
                    "class_id": class_id,
                    "subject_code": subject_code,
                    "subject_name": subject_name,
                }
                if doc.get("risk_source") is not None:
                    row["prediction_label"] = "External Factor" if doc.get("risk_source") == "external_factors" else "Academic Problem"
                if doc.get("gpa") is not None:
                    row["gpa"] = doc["gpa"]
                if doc.get("attendance") is not None:
                    row["attendance"] = doc["attendance"]
                if doc.get("lms_activity") is not None:
                    row["lms_activity"] = doc["lms_activity"]
                rows.append(row)
        return rows
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("", response_model=list[ClassResponse])
def list_classes(instructor_id: str, actor: dict = Depends(get_current_actor)):
    """List all active classes for an instructor."""
    try:
        db = get_db()
        cursor = db.classes.find({
            "instructor_id": instructor_id,
            "status": {"$ne": "archived"}
        }).sort("subject_code", 1)
        classes = []
        for doc in cursor:
            class_id = str(doc["_id"])
            count = db.enrollments.count_documents({"class_id": class_id})
            at_risk = db.enrollments.count_documents(
                {"class_id": class_id, "flagged_for_mentoring": True}
            )
            # Get section_code from first enrollment that has it
            enrollment = db.enrollments.find_one({"class_id": class_id, "section_code": {"$exists": True, "$ne": None}})
            section_code = enrollment.get("section_code") if enrollment else None
            classes.append(_doc_to_class_response(doc, count, at_risk, section_code))
        return classes
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}", response_model=ClassResponse)
def get_class(class_id: str, actor: dict = Depends(get_current_actor)):
    """Get a single class by id."""
    try:
        db = get_db()
        doc = _get_class_for_actor(db, class_id, actor)
        count = db.enrollments.count_documents({"class_id": class_id})
        at_risk = db.enrollments.count_documents(
            {"class_id": class_id, "flagged_for_mentoring": True}
        )
        # Get section_code from first enrollment that has it
        enrollment = db.enrollments.find_one({"class_id": class_id, "section_code": {"$exists": True, "$ne": None}})
        section_code = enrollment.get("section_code") if enrollment else None
        return _doc_to_class_response(doc, count, at_risk, section_code)
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.post("", response_model=ClassResponse, status_code=201)
def create_class(body: ClassCreate, actor: dict = Depends(get_current_actor)):
    """Create a new class (subject)."""
    try:
        db = get_db()
        section_code = re.sub(r"\s+", " ", body.section_code.strip())
        subject_code = re.sub(r"\s+", " ", body.subject_code.strip())
        subject_name = re.sub(r"\s+", " ", body.subject_name.strip())

        existing = db.classes.find_one({
            "instructor_id": body.instructor_id.strip(),
            "status": {"$ne": "archived"},
            "section_code": {"$regex": f"^{re.escape(section_code)}$", "$options": "i"},
            "subject_code": {"$regex": f"^{re.escape(subject_code)}$", "$options": "i"},
            "subject_name": {"$regex": f"^{re.escape(subject_name)}$", "$options": "i"},
        })
        if existing:
            raise HTTPException(status_code=400, detail="This class already exists.")

        doc = {
            "section_code": section_code,
            "subject_code": subject_code,
            "subject_name": subject_name,
            "instructor_id": body.instructor_id.strip(),
            "status": "active",
        }
        result = db.classes.insert_one(doc)
        doc["_id"] = result.inserted_id
        create_activity_log(
            db,
            actor_id=actor["id"],
            actor_name=actor.get("name", "User"),
            role=actor["role"],
            action="create_class",
            description=f"Created class {section_code} - {subject_code}: {subject_name}.",
            target_type="class",
            target_id=str(doc["_id"]),
            metadata={"section_code": section_code, "subject_code": subject_code, "subject_name": subject_name},
        )
        return _doc_to_class_response(doc, 0, 0, doc.get("section_code"))
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/archived/list")
def list_archived_classes(instructor_id: str, actor: dict = Depends(get_current_actor)):
    """List all archived classes for an instructor."""
    try:
        db = get_db()
        cursor = db.classes.find({
            "instructor_id": instructor_id,
            "status": "archived"
        }).sort("subject_code", 1)
        classes = []
        for doc in cursor:
            class_id = str(doc["_id"])
            count = db.enrollments.count_documents({"class_id": class_id})
            at_risk = db.enrollments.count_documents(
                {"class_id": class_id, "flagged_for_mentoring": True}
            )
            # Get section_code from first enrollment that has it
            enrollment = db.enrollments.find_one({"class_id": class_id, "section_code": {"$exists": True, "$ne": None}})
            section_code = enrollment.get("section_code") if enrollment else None
            classes.append(_doc_to_class_response(doc, count, at_risk, section_code))
        return classes
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.patch("/{class_id}/archive", response_model=ClassResponse)
def archive_class(class_id: str, actor: dict = Depends(get_current_actor)):
    """Archive a class and mark its linked enrollment data as archived."""
    try:
        db = get_db()
        doc = _get_class_for_actor(db, class_id, actor)
        if doc.get("status") == "archived":
            raise HTTPException(status_code=400, detail="Class is already archived")

        result = db.classes.update_one(
            {"_id": ObjectId(class_id)},
            {"$set": {"status": "archived"}}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Class not found")

        _set_related_enrollment_archive_state(db, class_id, archived=True)
        create_activity_log(
            db,
            actor_id=actor["id"],
            actor_name=actor.get("name", "User"),
            role=actor["role"],
            action="archive_class",
            description=f"Archived class {(doc.get('section_code') or '').strip()} {(doc.get('subject_code') or '').strip()} {(doc.get('subject_name') or '').strip()}".strip(),
            target_type="class",
            target_id=class_id,
        )
        
        doc = db.classes.find_one({"_id": ObjectId(class_id)})
        count = db.enrollments.count_documents({"class_id": class_id})
        at_risk = db.enrollments.count_documents(
            {"class_id": class_id, "flagged_for_mentoring": True}
        )
        # Get section_code from first enrollment that has it
        enrollment = db.enrollments.find_one({"class_id": class_id, "section_code": {"$exists": True, "$ne": None}})
        section_code = enrollment.get("section_code") if enrollment else None
        return _doc_to_class_response(doc, count, at_risk, section_code)
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.patch("/{class_id}/restore", response_model=ClassResponse)
def restore_class(class_id: str, actor: dict = Depends(get_current_actor)):
    """Restore an archived class together with its linked enrollment data."""
    try:
        db = get_db()
        doc = _get_class_for_actor(db, class_id, actor)
        if doc.get("status") != "archived":
            raise HTTPException(status_code=400, detail="Only archived classes can be restored")

        result = db.classes.update_one(
            {"_id": ObjectId(class_id)},
            {"$set": {"status": "active"}}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Class not found")

        _set_related_enrollment_archive_state(db, class_id, archived=False)
        create_activity_log(
            db,
            actor_id=actor["id"],
            actor_name=actor.get("name", "User"),
            role=actor["role"],
            action="restore_class",
            description=f"Restored class {(doc.get('section_code') or '').strip()} {(doc.get('subject_code') or '').strip()} {(doc.get('subject_name') or '').strip()}".strip(),
            target_type="class",
            target_id=class_id,
        )
        
        doc = db.classes.find_one({"_id": ObjectId(class_id)})
        count = db.enrollments.count_documents({"class_id": class_id})
        at_risk = db.enrollments.count_documents(
            {"class_id": class_id, "flagged_for_mentoring": True}
        )
        # Get section_code from first enrollment that has it
        enrollment = db.enrollments.find_one({"class_id": class_id, "section_code": {"$exists": True, "$ne": None}})
        section_code = enrollment.get("section_code") if enrollment else None
        return _doc_to_class_response(doc, count, at_risk, section_code)
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.delete("/{class_id}/permanent-delete", status_code=204)
def permanent_delete_class(class_id: str, actor: dict = Depends(get_current_actor)):
    """Permanently delete an archived class and all its data."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)
        class_obj_id = class_doc["_id"]
        if class_doc.get("status") != "archived":
            raise HTTPException(status_code=400, detail="Can only delete archived classes")
        
        # Delete all enrollments for this class
        db.enrollments.delete_many({"class_id": class_id})

        # Delete the class after related records are removed.
        db.classes.delete_one({"_id": class_obj_id})
        create_activity_log(
            db,
            actor_id=actor["id"],
            actor_name=actor.get("name", "User"),
            role=actor["role"],
            action="delete_class",
            description=f"Permanently deleted class {(class_doc.get('section_code') or '').strip()} {(class_doc.get('subject_code') or '').strip()} {(class_doc.get('subject_name') or '').strip()}".strip(),
            target_type="class",
            target_id=class_id,
        )
        
        return None
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}/students")
def list_class_students(class_id: str, actor: dict = Depends(get_current_actor)):
    """List students enrolled in a class with optional academic/risk/flag data."""
    try:
        db = get_db()
        if not ObjectId.is_valid(class_id):
            raise HTTPException(status_code=404, detail="Class not found")
        if not db.classes.find_one({"_id": ObjectId(class_id)}):
            raise HTTPException(status_code=404, detail="Class not found")
        # Sort by _id (insertion order) instead of student_email which may not exist
        cursor = db.enrollments.find({"class_id": class_id}).sort("_id", 1)
        out = []
        ai_fields = [
            "previous_gpa",
            "previous_midterm_class_standing",
            "previous_midterm_laboratory",
            "previous_midterm_major_output",
            "previous_midterm_grade",
            "previous_final_class_standing",
            "previous_final_laboratory",
            "previous_final_major_output",
            "previous_final_grade",
            "previous_failed_flag",
            "previous_passed_flag",
            "historical_grade_average",
            "historical_failure_count",
            "failed_subject_count",
            "attendance",
            "attendance_overall",
            "self_reported_attendance",
            "class_standing",
            "laboratory",
            "major_output",
            "midterm_grade",
            "final_class_standing",
            "final_laboratory",
            "final_major_output",
            "final_grade",
            "overall_grade",
            "model_features",
            "model_profile",
            "risk_source",
            "risk_source_label",
            "risk_drivers",
            "academic_risk_drivers",
            "external_risk_drivers",
            "midterm_topic_difficulties",
            "hardest_midterm_topics",
            "top_contributing_signals",
            "received_academic_support",
            "on_probation_status",
            "has_subject_grade_2_5",
            "gwa_2_5_or_below",
            "low_midterm_academic_performance",
            "difficulty_catching_up_instructions",
            "difficulty_understanding_lectures",
            "struggles_specific_subjects",
            "weak_study_habits_time_management",
            "low_motivation_engagement",
            "poor_comprehension_writing_skills",
            "financial_difficulties",
            "physical_health_concerns",
            "family_issues",
            "part_time_work_affecting_studies",
            "mental_health_concerns",
        ]
        for doc in cursor:
            student_email, student_name, student_id, student_key = _get_enrollment_identity(doc)
            resolved_name = _enrich_student_name(db, student_email, student_id, student_name)
            row = {
                "student_email": student_email or student_key,
                "student_name": resolved_name or student_name or None,
                "student_id": student_id or None,
            }
            if doc.get("risk_source") is not None:
                row["prediction_label"] = "External Factor" if doc.get("risk_source") == "external_factors" else "Academic Problem"
            if doc.get("gpa") is not None:
                row["gpa"] = doc["gpa"]
            if doc.get("attendance") is not None:
                row["attendance"] = doc["attendance"]
            if doc.get("lms_activity") is not None:
                row["lms_activity"] = doc["lms_activity"]
            if doc.get("flagged_for_mentoring") is not None:
                row["flagged_for_mentoring"] = doc["flagged_for_mentoring"]
            if doc.get("referral_note") is not None:
                row["referral_note"] = doc["referral_note"]
            if doc.get("assigned_amu_staff_id") is not None:
                row["assigned_amu_staff_id"] = doc["assigned_amu_staff_id"]
            if doc.get("assigned_amu_staff_name") is not None:
                row["assigned_amu_staff_name"] = doc["assigned_amu_staff_name"]
            if doc.get("assigned_amu_staff_college") is not None:
                row["assigned_amu_staff_college"] = doc["assigned_amu_staff_college"]
            for field in ai_fields:
                if doc.get(field) is not None:
                    row[field] = doc[field]
            out.append(row)
        return out
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}/risk-summary")
def get_class_risk_summary(class_id: str, actor: dict = Depends(get_current_actor)):
    """Class-level outcome summary for academic vs external prediction labels."""
    try:
        db = get_db()
        if not ObjectId.is_valid(class_id):
            raise HTTPException(status_code=404, detail="Class not found")
        if not db.classes.find_one({"_id": ObjectId(class_id)}):
            raise HTTPException(status_code=404, detail="Class not found")
        cursor = list(db.enrollments.find({"class_id": class_id}))
        total = len(cursor)
        academic_count = sum(1 for d in cursor if d.get("risk_source") == "academic")
        external_count = sum(1 for d in cursor if d.get("risk_source") == "external_factors")
        referred_count = sum(1 for d in cursor if d.get("flagged_for_mentoring"))
        outcome_list = []
        for d in cursor:
            if d.get("risk_source") not in ("academic", "external_factors"):
                continue
            student_email, student_name, student_id, student_key = _get_enrollment_identity(d)
            resolved_name = _enrich_student_name(db, student_email, student_id, student_name)
            outcome_list.append({
                "student_email": student_email or student_key,
                "student_name": resolved_name or None,
                "student_id": student_id or None,
                "prediction_label": "External Factor" if d.get("risk_source") == "external_factors" else "Academic Problem",
            })
        return {
            "total": total,
            "academic_count": academic_count,
            "external_count": external_count,
            "referred_count": referred_count,
            "outcome_list": outcome_list,
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}/roster")
def get_class_roster(class_id: str, actor: dict = Depends(get_current_actor)):
    """Fetch student roster for a class with names and IDs."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)
        
        section_code = None
        students = []
        
        # Get all enrollments for this class
        enrollments = list(db.enrollments.find({"class_id": class_id}))
        
        for enrollment in enrollments:
            student_email, student_name, student_id, student_key = _get_enrollment_identity(enrollment)
            resolved_name = _enrich_student_name(db, student_email, student_id, student_name)

            student_info = {
                "email": student_email or student_key,
                "name": resolved_name or student_name or "Unknown",
                "student_id": student_id or "",
            }
            # Capture section code from first enrollment that has it
            if not section_code and enrollment.get("section_code"):
                section_code = enrollment.get("section_code")
            students.append(student_info)
        
        return {
            "class_id": class_id,
            "section_code": section_code,
            "students": students
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}/grades")
def get_class_grades_with_analytics(class_id: str, actor: dict = Depends(get_current_actor)):
    """Fetch all student grades with class analytics."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)
        
        enrollments = list(db.enrollments.find({"class_id": class_id}))
        
        # Build student grades list
        students = []
        grade_values = []
        midterm_grades = []
        final_grades = []
        pass_count = 0
        fail_count = 0
        has_breakdown_scores = any(bool(enrollment.get("grades_breakdown")) for enrollment in enrollments)
        score_columns, score_column_terms = _collect_class_score_column_metadata(enrollments)
        activity_title_mappings = class_doc.get("activity_title_mappings") or {}
        
        for enrollment in enrollments:
            student_email, student_name, student_id, student_key = _get_enrollment_identity(enrollment)
            resolved_name = _enrich_student_name(db, student_email, student_id, student_name)

            stored_scores = dict(enrollment.get("grades_breakdown") or {})
            stored_column_order = list(enrollment.get("grades_column_order") or list(stored_scores.keys()))
            stored_terms = dict(enrollment.get("grades_column_terms") or {})

            stored_score_aliases = _build_raw_score_column_aliases(
                stored_column_order,
                set(),
                alias_all_section_columns=False,
            )

            # Return only raw Class Standing/Lab/MO component scores.
            scores = {}
            column_order = []
            for source_col in stored_column_order:
                if source_col not in stored_scores:
                    continue
                value = stored_scores[source_col]
                score_col = stored_score_aliases.get(source_col, source_col)
                if not _should_include_score_column(score_col, _normalize_cell(value), value):
                    continue
                scores[score_col] = value
                inferred_term = (
                    _normalize_score_term(stored_terms.get(source_col))
                    or _normalize_score_term(stored_terms.get(score_col))
                    or ("final" if _is_final_term_hint_column(source_col) or _is_final_term_hint_column(score_col) else None)
                    or "midterm"
                )
                score_column_terms.setdefault(score_col, inferred_term)
                if score_col not in column_order:
                    column_order.append(score_col)

            # Backward compatibility: merge scalar legacy fields into score payload.
            if not has_breakdown_scores:
                legacy_scores = _legacy_grade_score_map(enrollment)
                legacy_aliases = _build_raw_score_column_aliases(
                    list(legacy_scores.keys()),
                    set(),
                    alias_all_section_columns=True,
                )
                for legacy_col, legacy_value in legacy_scores.items():
                    score_col = legacy_aliases.get(legacy_col, legacy_col)
                    if score_col not in scores:
                        scores[score_col] = legacy_value
                    legacy_term = "final" if legacy_col.startswith("final_") else "midterm"
                    score_column_terms.setdefault(score_col, legacy_term)
                    if score_col not in column_order:
                        column_order.append(score_col)

            # Keep student-level order grouped by CS -> LAB -> MO, each in numeric item order.
            column_order = sorted(column_order, key=_raw_score_column_sort_key)

            student = {
                "id": str(enrollment.get("_id", "")),
                "email": student_email or "",
                "name": resolved_name or None,
                "id_number": student_id or enrollment.get("id_number", ""),
                "class_standing": enrollment.get("class_standing"),
                "laboratory": enrollment.get("laboratory"),
                "major_output": enrollment.get("major_output"),
                "final_class_standing": enrollment.get("final_class_standing"),
                "final_laboratory": enrollment.get("final_laboratory"),
                "final_major_output": enrollment.get("final_major_output"),
                "summary": enrollment.get("summary"),
                "midterm_grade": enrollment.get("midterm_grade"),
                "final_grade": enrollment.get("final_grade"),
                "midterm_weighted": enrollment.get("midterm_weighted"),
                "final_weighted": enrollment.get("final_weighted"),
                "overall_grade": enrollment.get("overall_grade"),
                "scores": scores,
                "section_code": enrollment.get("section_code", class_doc.get("section_code", "")),
                "subject_code": enrollment.get("subject_code", class_doc.get("subject_code", "")),
                "class_time": enrollment.get("class_time", class_doc.get("class_time", "")),
                "gpa": enrollment.get("gpa"),
                "risk": enrollment.get("risk"),
            }

            for numeric_key in (
                "class_standing",
                "laboratory",
                "major_output",
                "final_class_standing",
                "final_laboratory",
                "final_major_output",
                "summary",
                "midterm_grade",
                "final_grade",
                "midterm_weighted",
                "final_weighted",
                "overall_grade",
                "gpa",
            ):
                student[numeric_key] = _to_float(student.get(numeric_key))
            
            students.append(student)
            
            # Collect grades for analytics
            if isinstance(student["gpa"], (int, float)):
                grade_values.append(student["gpa"])
            if isinstance(student["midterm_grade"], (int, float)):
                midterm_grades.append(student["midterm_grade"])
            if isinstance(student["final_grade"], (int, float)):
                final_grades.append(student["final_grade"])
            
            # Count pass/fail for both 1.0-5.0 and 0-100 grade scales.
            if isinstance(student["gpa"], (int, float)):
                if student["gpa"] <= 5:
                    if student["gpa"] <= 3.0:
                        pass_count += 1
                    else:
                        fail_count += 1
                else:
                    if student["gpa"] >= 60:
                        pass_count += 1
                    else:
                        fail_count += 1
        
        analytics = {
            "total_students": len(students),
            "gpa_average": round(sum(grade_values) / len(grade_values), 2) if grade_values else 0,
            "gpa_highest": max(grade_values) if grade_values else 0,
            "gpa_lowest": min(grade_values) if grade_values else 0,
            "midterm_average": round(sum(midterm_grades) / len(midterm_grades), 2) if midterm_grades else 0,
            "final_average": round(sum(final_grades) / len(final_grades), 2) if final_grades else 0,
            "pass_rate": round(pass_count / len(students) * 100, 2) if students else 0,
            "fail_rate": round(fail_count / len(students) * 100, 2) if students else 0,
        }
        
        return {
            "class": {
                "id": class_id,
                "subject_code": class_doc.get("subject_code", ""),
                "subject_name": class_doc.get("subject_name", ""),
            },
            "score_columns": score_columns,
            "score_column_terms": score_column_terms,
            "activity_title_mappings": activity_title_mappings,
            "activity_title_mapping_count": len(activity_title_mappings) if isinstance(activity_title_mappings, dict) else 0,
            "students": students,
            "analytics": analytics,
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.get("/{class_id}/attendance")
def get_class_attendance_with_analytics(class_id: str, actor: dict = Depends(get_current_actor)):
    """Fetch all student attendance records with class analytics."""
    try:
        db = get_db()
        if not ObjectId.is_valid(class_id):
            raise HTTPException(status_code=404, detail="Class not found")
        
        class_doc = db.classes.find_one({"_id": ObjectId(class_id)})
        if not class_doc:
            raise HTTPException(status_code=404, detail="Class not found")
        
        enrollments = list(db.enrollments.find({"class_id": class_id}))
        
        # Build student attendance list
        students = []
        overall_attendance_values = []
        present_days_values = []
        absent_days_values = []
        monthly_attendance = {month: [] for month in ['january', 'february', 'march', 'april', 'may', 'june', 
                                                       'july', 'august', 'september', 'october', 'november', 'december']}
        
        # Detect if using daily or monthly format
        is_daily_format = False
        for enrollment in enrollments:
            if enrollment.get('attendance_present_days') is not None or enrollment.get('attendance_absent_days') is not None:
                is_daily_format = True
                break
        
        for enrollment in enrollments:
            student_email, student_name, student_id, _ = _get_enrollment_identity(enrollment)
            resolved_name = _enrich_student_name(db, student_email, student_id, student_name)
            student = {
                "id": str(enrollment.get("_id", "")),
                "email": student_email,
                "name": resolved_name or None,
                "id_number": student_id or enrollment.get("id_number", ""),
                "section_code": enrollment.get("section_code", class_doc.get("section_code", "")),
                "subject_code": enrollment.get("subject_code", class_doc.get("subject_code", "")),
            }

            # Handle daily attendance format
            if is_daily_format:
                present_days = enrollment.get("attendance_present_days")
                absent_days = enrollment.get("attendance_absent_days")
                if present_days is not None:
                    student["present_days"] = present_days
                if absent_days is not None:
                    student["absent_days"] = absent_days
                if present_days is not None or absent_days is not None:
                    present_days_values.append(present_days or 0)
                    absent_days_values.append(absent_days or 0)
            else:
                # Handle monthly attendance format
                monthly_data = {}
                for month in ['january', 'february', 'march', 'april', 'may', 'june', 
                             'july', 'august', 'september', 'october', 'november', 'december']:
                    attendance_key = f'attendance_{month}'
                    value = enrollment.get(attendance_key)
                    monthly_data[month] = value
                    if value is not None:
                        monthly_attendance[month].append(value)
                
                student["attendance"] = monthly_data
            
            # Get overall attendance
            overall = enrollment.get("attendance_overall")
            if overall is not None:
                student["overall_attendance"] = overall
                overall_attendance_values.append(overall)
            
            students.append(student)
        
        # Calculate analytics
        analytics = {
            "total_students": len(students),
            "attendance_format": "daily" if is_daily_format else "monthly",
        }
        
        if is_daily_format:
            # Daily attendance analytics
            total_present = sum(present_days_values)
            total_absent = sum(absent_days_values)
            total_days = total_present + total_absent
            analytics.update({
                "total_present_days": total_present,
                "total_absent_days": total_absent,
                "total_attendance_days": total_days,
                "average_present_per_student": round(total_present / len(students), 2) if students else 0,
                "average_absent_per_student": round(total_absent / len(students), 2) if students else 0,
            })
            
            # High absenteeism flag (students with more absences than presence)
            high_absenteeism = sum(1 for s in students if s.get("absent_days", 0) > s.get("present_days", 0))
            analytics["high_absenteeism_count"] = high_absenteeism
            analytics["high_absenteeism_percentage"] = round(high_absenteeism / len(students) * 100, 2) if students else 0
        else:
            # Monthly attendance analytics
            analytics.update({
                "overall_average": round(sum(overall_attendance_values) / len(overall_attendance_values), 2) if overall_attendance_values else 0,
                "overall_highest": max(overall_attendance_values) if overall_attendance_values else 0,
                "overall_lowest": min(overall_attendance_values) if overall_attendance_values else 0,
            })
            
            # Add monthly averages
            for month, values in monthly_attendance.items():
                if values:
                    analytics[f"{month}_average"] = round(sum(values) / len(values), 2)
                else:
                    analytics[f"{month}_average"] = 0
            
            # Calculate students with low attendance (<75%)
            low_attendance_count = sum(1 for s in students if s.get("overall_attendance") and s["overall_attendance"] < 75)
            analytics["low_attendance_count"] = low_attendance_count
            analytics["low_attendance_percentage"] = round(low_attendance_count / len(students) * 100, 2) if students else 0
        
        return {
            "class": {
                "id": class_id,
                "subject_code": class_doc.get("subject_code", ""),
                "subject_name": class_doc.get("subject_name", ""),
            },
            "students": students,
            "analytics": analytics,
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


# Single-email add endpoint removed. Use batch add or classlist upload instead.


@router.post("/{class_id}/students/batch", status_code=201)
def batch_add_students_to_class(class_id: str, body: BatchAddStudentsRequest, actor: dict = Depends(get_current_actor)):
    """Add multiple students to a class by email list."""
    try:
        db = get_db()
        if not ObjectId.is_valid(class_id):
            raise HTTPException(status_code=404, detail="Class not found")
        if not db.classes.find_one({"_id": ObjectId(class_id)}):
            raise HTTPException(status_code=404, detail="Class not found")
        added = 0
        skipped = 0
        for raw in body.emails:
            email = raw.strip().lower()
            if not email:
                continue
            existing = db.enrollments.find_one({"class_id": class_id, "student_email": email})
            if existing:
                skipped += 1
                continue
            db.enrollments.insert_one({"class_id": class_id, "student_email": email})
            added += 1
        return {"message": "Batch add complete.", "added": added, "skipped": skipped}
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.patch("/{class_id}/students/{student_identifier:path}")
def update_enrollment(class_id: str, student_identifier: str, body: UpdateEnrollmentRequest, actor: dict = Depends(get_current_actor)):
    """Update academic indicators, risk, or flagged_for_mentoring for a student in the class."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)

        identifier = _normalize_cell(student_identifier).strip()
        identifier_lower = identifier.lower()
        match_filter = {
            "class_id": class_id,
            "$or": [
                {"student_email": identifier_lower},
                {"student_id": identifier},
                {"student_id": identifier_lower},
            ],
        }
        doc = db.enrollments.find_one(match_filter)
        if not doc:
            raise HTTPException(status_code=404, detail="Student not enrolled in this class.")
        payload = body.model_dump(exclude_unset=True)
        student_email = (doc.get("student_email") or "").strip().lower()
        student_id = _normalize_cell(doc.get("student_id"))
        student_label = student_email or student_id or identifier
        if not payload:
            return {"message": "No updates.", "student_email": student_email or None, "student_id": student_id or None}
        if payload.get("flagged_for_mentoring") is True:
            # Debug: log the referral_reasons being received
            if "referral_reasons" in payload:
                print(f"DEBUG: Referral reasons received: {payload['referral_reasons']}")
            assigned_amu_staff_id = str(payload.get("assigned_amu_staff_id") or "").strip()
            assigned_amu_staff_name = str(payload.get("assigned_amu_staff_name") or "").strip()
            assigned_amu_staff_college = str(payload.get("assigned_amu_staff_college") or "").strip()
            if not assigned_amu_staff_id or not assigned_amu_staff_name:
                inferred_doc = {**doc, **payload}
                auto_reasons = _build_auto_referral_reasons(inferred_doc)
                if any(auto_reasons.values()):
                    instructor_doc = None
                    instructor_id = class_doc.get("instructor_id")
                    if instructor_id and ObjectId.is_valid(instructor_id):
                        instructor_doc = db.instructor.find_one({"_id": ObjectId(instructor_id)})
                    auto_staff = _find_amu_staff_for_college(db, (instructor_doc or {}).get("college"))
                    if auto_staff:
                        assigned_amu_staff_id = str(auto_staff["_id"])
                        assigned_amu_staff_name = _normalize_cell(auto_staff.get("name"))
                        assigned_amu_staff_college = _normalize_cell(auto_staff.get("college"))
                        payload.setdefault("referral_reasons", auto_reasons)
            if not assigned_amu_staff_id or not assigned_amu_staff_name:
                raise HTTPException(
                    status_code=400,
                    detail="Please choose an AMU staff member before sending the referral.",
                )
            payload["assigned_amu_staff_id"] = assigned_amu_staff_id
            payload["assigned_amu_staff_name"] = assigned_amu_staff_name
            payload["assigned_amu_staff_college"] = assigned_amu_staff_college or None
        db.enrollments.update_one(
            {"_id": doc["_id"]},
            {"$set": payload},
        )
        refreshed_doc = db.enrollments.find_one({"_id": doc["_id"]})
        if refreshed_doc:
            _apply_automatic_referral(db, class_doc, refreshed_doc)
        # When instructor flags a student for mentoring, set referred_at and notify AMU staff
        if payload.get("flagged_for_mentoring") is True:
            now = datetime.now(timezone.utc)
            db.enrollments.update_one(
                {"_id": doc["_id"]},
                {"$set": {"referred_at": now}},
            )
            instructor_name = "Instructor"
            if class_doc and class_doc.get("instructor_id"):
                inst_doc = db.instructor.find_one({"_id": ObjectId(class_doc["instructor_id"])})
                if inst_doc and inst_doc.get("name"):
                    instructor_name = inst_doc["name"]
            subject_code = (class_doc or {}).get("subject_code") or "Class"
            subject_name = (class_doc or {}).get("subject_name") or ""
            class_label = f"{subject_code}" + (f": {subject_name}" if subject_name else "")
            assigned_label = payload.get("assigned_amu_staff_name") or "an AMU staff member"
            assigned_college = payload.get("assigned_amu_staff_college")
            if assigned_college:
                assigned_label = f"{assigned_label} - {assigned_college}"
            create_notification(
                db,
                role="amu-staff",
                recipient_user_id=payload.get("assigned_amu_staff_id"),
                title="Student flagged by instructor",
                body=f"Student {student_label} was flagged by {instructor_name} ({class_label}) and assigned to {assigned_label}. Please review the case.",
                type="alert",
            )
        return {"message": "Enrollment updated.", "student_email": student_email or None, "student_id": student_id or None}
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.post("/{class_id}/students/{student_email:path}/predict-risk")
def predict_enrollment_risk(class_id: str, student_email: str, actor: dict = Depends(get_current_actor)):
    """Run the XGBoost student-risk model for one enrolled student."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)

        email = student_email.strip().lower()
        doc = db.enrollments.find_one({"class_id": class_id, "student_email": email})
        if not doc:
            raise HTTPException(status_code=404, detail="Student not enrolled in this class.")

        try:
            result = predict_student_risk(doc)
        except FileNotFoundError as exc:
            raise HTTPException(status_code=500, detail=str(exc)) from exc
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Prediction failed: {exc}") from exc

        db.enrollments.update_one(
            {"class_id": class_id, "student_email": email},
            {
                "$set": {
                    "model_features": result["features"],
                    "risk_source": result.get("risk_source"),
                    "risk_source_label": result.get("risk_source_label"),
                    "risk_drivers": result.get("risk_drivers"),
                    "academic_risk_drivers": result.get("academic_risk_drivers"),
                    "external_risk_drivers": result.get("external_risk_drivers"),
                    "model_profile": result.get("model_profile"),
                    "midterm_topic_difficulties": result.get("midterm_topic_difficulties"),
                    "hardest_midterm_topics": result.get("hardest_midterm_topics"),
                    "top_contributing_signals": result.get("top_contributing_signals"),
                }
            ,
                "$unset": {
                    "risk": "",
                    "risk_prediction": "",
                    "risk_probability": "",
                    "risk_probability_percent": "",
                }
            },
        )
        refreshed = db.enrollments.find_one({"class_id": class_id, "student_email": email})
        if refreshed:
            _apply_automatic_referral(db, class_doc, refreshed)

        return {
            "class_id": class_id,
            "student_email": email,
            **result,
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")


@router.post("/{class_id}/upload-needs-assessment", status_code=201)
async def upload_needs_assessment_file(
    class_id: str,
    actor: dict = Depends(get_current_actor),
    files: List[UploadFile] = File(..., description="CSV or XLSX files with needs-assessment columns"),
):
    """Bulk upload needs-assessment fields for students in a class."""
    if not ObjectId.is_valid(class_id):
        raise HTTPException(status_code=404, detail="Class not found")

    try:
        db = get_db()
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")

    class_doc = _get_class_for_actor(db, class_id, actor)

    saved_files = []
    updated = 0
    not_enrolled = []
    missing_identifiers = 0

    for upload in files:
        ext = os.path.splitext(upload.filename or "")[1].lower()
        if ext not in [".csv", ".xlsx"]:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Please upload CSV or XLSX files only.")

        unique_name = f"{class_id}_needs_assessment_{uuid.uuid4().hex}{ext}"
        dest = UPLOAD_DIR / unique_name
        with dest.open("wb") as buffer:
            shutil.copyfileobj(upload.file, buffer)
        saved_files.append(str(dest.name))

        try:
            rows = _parse_file_to_rows(dest, preferred_type="gradesheet")
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse {upload.filename}: {exc}")

        if not rows:
            raise HTTPException(status_code=400, detail=f"{upload.filename} did not contain any readable rows.")

        keys = list(rows[0].keys())
        for row in rows:
            enrollment, lookup_identifier, _email, _name, _id = _find_matching_enrollment(db, class_id, row, keys)
            if not enrollment:
                if not _row_identifier_label(row, keys):
                    missing_identifiers += 1
                if lookup_identifier:
                    not_enrolled.append(lookup_identifier)
                continue

            update_data = _build_ai_input_update_data(row, keys)
            if not update_data:
                continue

            db.enrollments.update_one({"_id": enrollment["_id"]}, {"$set": update_data})
            refreshed = db.enrollments.find_one({"_id": enrollment["_id"]})
            if refreshed:
                _apply_automatic_referral(db, class_doc, refreshed)
            updated += 1

    return {
        "message": "Needs Assessment file(s) uploaded successfully.",
        "files": saved_files,
        "updated": updated,
        "not_enrolled": sorted(set(not_enrolled))[:25],
        "missing_identifiers": missing_identifiers,
    }


@router.post("/{class_id}/upload-activity-titles", status_code=201)
async def upload_activity_title_mapping_file(
    class_id: str,
    actor: dict = Depends(get_current_actor),
    files: List[UploadFile] = File(..., description="CSV or XLSX files with activity-title mappings"),
):
    """Upload a class-level activity-title mapping for score columns."""
    if not ObjectId.is_valid(class_id):
        raise HTTPException(status_code=404, detail="Class not found")

    try:
        db = get_db()
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")

    class_doc = _get_class_for_actor(db, class_id, actor)
    enrollments = list(db.enrollments.find({"class_id": class_id}))
    score_columns, score_column_terms = _collect_class_score_column_metadata(enrollments)
    if not score_columns:
        raise HTTPException(
            status_code=400,
            detail="Upload a current gradesheet first before importing activity titles.",
        )

    existing_mappings = class_doc.get("activity_title_mappings") or {}
    saved_files = []
    merged_mappings = dict(existing_mappings) if isinstance(existing_mappings, dict) else {}
    unmatched_rows = 0
    matched_rows = 0

    for upload in files:
        ext = os.path.splitext(upload.filename or "")[1].lower()
        if ext not in [".csv", ".xlsx"]:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Please upload CSV or XLSX files only.")

        unique_name = f"{class_id}_activity_titles_{uuid.uuid4().hex}{ext}"
        dest = UPLOAD_DIR / unique_name
        with dest.open("wb") as buffer:
            shutil.copyfileobj(upload.file, buffer)
        saved_files.append(str(dest.name))

        try:
            rows = _parse_file_to_rows(dest, preferred_type="gradesheet")
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse {upload.filename}: {exc}")

        if not rows:
            raise HTTPException(status_code=400, detail=f"{upload.filename} did not contain any readable rows.")

        file_mappings, file_unmatched_rows = _build_activity_title_mapping_payload(
            rows,
            score_columns,
            score_column_terms,
            merged_mappings,
        )
        merged_mappings = file_mappings
        unmatched_rows += file_unmatched_rows

    matched_rows = len(merged_mappings)
    if not merged_mappings:
        raise HTTPException(
            status_code=400,
            detail=(
                "No activity titles matched your class score columns. "
                "Use either 'activity column' + 'activity title', or 'term' + 'component' + 'item' + 'activity title'."
            ),
        )

    updated_at = datetime.now(timezone.utc)
    db.classes.update_one(
        {"_id": ObjectId(class_id)},
        {
            "$set": {
                "activity_title_mappings": merged_mappings,
                "activity_title_mapping_files": saved_files,
                "activity_title_mapping_updated_at": updated_at,
            }
        },
    )

    for enrollment in enrollments:
        topic_payload = _extract_topic_difficulty(
            {
                **enrollment,
                "activity_title_mappings": merged_mappings,
            },
            enrollment.get("model_features") or {},
        )
        db.enrollments.update_one(
            {"_id": enrollment["_id"]},
            {
                "$set": {
                    "activity_title_mappings": merged_mappings,
                    "midterm_topic_difficulties": topic_payload.get("midterm_topic_difficulties") or [],
                    "hardest_midterm_topics": topic_payload.get("hardest_midterm_topics") or [],
                }
            },
        )

    return {
        "message": "Activity title mapping uploaded successfully.",
        "files": saved_files,
        "matched": matched_rows,
        "unmatched_rows": unmatched_rows,
        "activity_title_mappings": merged_mappings,
    }


@router.post("/{class_id}/predict-risk", status_code=200)
def predict_class_risk(class_id: str, actor: dict = Depends(get_current_actor)):
    """Run risk prediction for all students enrolled in a class."""
    try:
        db = get_db()
        class_doc = _get_class_for_actor(db, class_id, actor)

        cursor = list(db.enrollments.find({"class_id": class_id}).sort("_id", 1))
        predicted = 0
        skipped = []
        results = []

        for doc in cursor:
            email = (doc.get("student_email") or "").strip().lower()
            student_key = email or doc.get("student_id") or doc.get("student_name") or str(doc.get("_id"))

            try:
                result = predict_student_risk(doc)
            except FileNotFoundError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc
            except Exception as exc:
                skipped.append(student_key)
                continue

            db.enrollments.update_one(
                {"_id": doc["_id"]},
                {
                    "$set": {
                        "model_features": result["features"],
                        "risk_source": result.get("risk_source"),
                        "risk_source_label": result.get("risk_source_label"),
                        "risk_drivers": result.get("risk_drivers"),
                        "academic_risk_drivers": result.get("academic_risk_drivers"),
                        "external_risk_drivers": result.get("external_risk_drivers"),
                        "model_profile": result.get("model_profile"),
                        "midterm_topic_difficulties": result.get("midterm_topic_difficulties"),
                        "hardest_midterm_topics": result.get("hardest_midterm_topics"),
                        "top_contributing_signals": result.get("top_contributing_signals"),
                    }
                ,
                    "$unset": {
                        "risk": "",
                        "risk_prediction": "",
                        "risk_probability": "",
                        "risk_probability_percent": "",
                    }
                },
            )
            refreshed = db.enrollments.find_one({"_id": doc["_id"]})
            if refreshed:
                _apply_automatic_referral(db, class_doc, refreshed)
            predicted += 1
            results.append({
                "student_email": email or None,
                "student_key": student_key,
                "prediction_label": "External Factor" if result.get("risk_source") == "external_factors" else "Academic Problem",
                "probability_percent": result["probability_percent"],
            })

        return {
            "class_id": class_id,
            "predicted": predicted,
            "skipped": skipped,
            "results": results,
        }
    except ServerSelectionTimeoutError:
        raise HTTPException(status_code=503, detail="Database unavailable.")
