from docx import Document
from pathlib import Path

doc_path = Path("backend/uploads/monthly_class_attendance_with_real_names.docx")
doc = Document(str(doc_path))

print("=" * 80)
print("ATTENDANCE SHEET ANALYSIS")
print("=" * 80)

if not doc.tables:
    print("No tables found in document")
else:
    table = doc.tables[0]
    print(f"\nTable Found: {len(table.rows)} rows, {len(table.columns)} columns")
    
    # Extract headers
    header_row = table.rows[0]
    headers = [cell.text.strip() for cell in header_row.cells]
    print(f"\nHeaders ({len(headers)} columns):")
    for i, h in enumerate(headers, 1):
        print(f"  {i}. {h}")
    
    # Extract first 5 data rows for preview
    print(f"\nData Sample (first 5 rows):")
    print("-" * 80)
    for row_idx, row in enumerate(table.rows[1:6], 1):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"Row {row_idx}: {cells}")
    
    # Extract all data rows for analysis
    print(f"\n\nDetailed Data Structure:")
    print("-" * 80)
    all_rows = []
    for row in table.rows[1:]:
        row_data = {}
        for col_idx, cell in enumerate(row.cells):
            header = headers[col_idx] if col_idx < len(headers) else f"Col_{col_idx}"
            row_data[header] = cell.text.strip()
        all_rows.append(row_data)
    
    # Print first 3 rows in detail
    for idx, row in enumerate(all_rows[:3], 1):
        print(f"\nRow {idx}:")
        for key, val in row.items():
            print(f"  {key}: {val}")
    
    print(f"\n... ({len(all_rows) - 3} more rows)")
    print(f"\nTotal data rows: {len(all_rows)}")
    
    # Analyze month columns
    print(f"\n\nMonth Columns Detected:")
    print("-" * 80)
    months = ['january', 'february', 'march', 'april', 'may', 'june', 
              'july', 'august', 'september', 'october', 'november', 'december']
    found_months = []
    for header in headers:
        header_lower = header.lower()
        for month in months:
            if month in header_lower:
                found_months.append((header, month))
                break
    
    for header, month in found_months:
        print(f"  ✓ {header} → {month}")
    
    print(f"\nTotal month columns: {len(found_months)}")
    
    # Analyze ID/Name columns
    print(f"\n\nIdentifier Columns:")
    print("-" * 80)
    for header in headers:
        header_lower = header.lower()
        if any(x in header_lower for x in ['id', 'no.', 'no', 'number', 'sid']):
            print(f"  ✓ ID Column: {header}")
        if any(x in header_lower for x in ['name', 'student name', 'full name']):
            print(f"  ✓ Name Column: {header}")
        if 'email' in header_lower:
            print(f"  ✓ Email Column: {header}")
