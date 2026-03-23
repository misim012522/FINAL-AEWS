from docx import Document
from pathlib import Path

# Simulate the backend parsing logic
def _find_column(keys, keywords):
    """Find first key containing any of the keywords."""
    for kw in keywords:
        for key in keys:
            if kw in key.lower():
                return key
    return None

def _is_daily_attendance_format(keys):
    """Detect if file is daily attendance format (checkmarks/signatures instead of percentages)."""
    attendance_keywords = ['date', 'signature', 'face-to-face', 'f2f', 'synchronous', 'asynchronous', 'async']
    has_attendance_markers = any(_find_column(keys, [kw]) for kw in attendance_keywords)
    
    # Check if there are NO month keywords (not monthly format)
    months = ['january', 'february', 'march', 'april', 'may', 'june',
              'july', 'august', 'september', 'october', 'november', 'december']
    has_months = any(_find_column(keys, [month]) for month in months)
    
    return has_attendance_markers and not has_months

def _parse_daily_attendance(row, keys, id_col, name_col):
    """Parse daily attendance from checkmarks."""
    marked_columns = []
    present_days = 0
    total_attendance_cols = 0
    
    name_keywords = ['name of students', 'name of student', 'name', 'student name', 'student_name']
    id_keywords = ['no.', 'no', 'number', 'id', 'student id', 'id number']
    
    for col_name in keys:
        # Skip ID and Name columns
        is_name = any(kw in col_name.lower() for kw in name_keywords)
        is_id = any(kw in col_name.lower() for kw in id_keywords)
        if is_name or is_id or col_name == id_col or col_name == name_col:
            continue
        
        cell_value = row.get(col_name, '').strip()
        if cell_value or col_name in [k for k in keys if k not in ['name of students', id_col, name_col]]:
            total_attendance_cols += 1
            if '✓' in cell_value or cell_value.lower() in ['present', 'p', 'yes']:
                present_days += 1
                marked_columns.append(col_name)
    
    absent_days = total_attendance_cols - present_days if total_attendance_cols > 0 else 0
    attendance_pct = (present_days / total_attendance_cols * 100) if total_attendance_cols > 0 else 0
    
    return present_days, absent_days, round(attendance_pct, 2), marked_columns

# Load the actual document
doc_path = Path("backend/uploads/monthly_class_attendance_with_real_names.docx")
doc = Document(str(doc_path))

if not doc.tables:
    print("No tables found")
else:
    table = doc.tables[0]
    
    # Parse header - skip first 3 rows (main headers, sub-headers, signature row)
    header_row = table.rows[3]  # Skip to 4th row which should be first data row for reference
    
    # Actually, let's use row 0 as headers
    header_row = table.rows[0]
    keys = [cell.text.strip().lower() for cell in header_row.cells]
    
    print("DETECTED FORMAT TEST")
    print("=" * 80)
    print(f"Keys: {keys[:3]}... (total {len(keys)} columns)")
    
    is_daily = _is_daily_attendance_format(keys)
    print(f"\nIs Daily Format: {is_daily}")
    
    if is_daily:
        print("\n✓ Daily attendance format detected!\n")
        
        id_col = _find_column(keys, ['no.', 'no', 'number'])
        name_col = _find_column(keys, ['name of students', 'name'])
        
        print(f"ID Column: {id_col}")
        print(f"Name Column: {name_col}")
        
        # Parse first 5 students
        print("\nParsed Student Data:")
        print("-" * 80)
        
        for row_idx in range(3, min(8, len(table.rows))):  # Skip to actual data rows
            row = table.rows[row_idx]
            row_dict = {keys[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
            
            if not row_dict.get(id_col):
                continue
            
            present, absent, pct, marked = _parse_daily_attendance(row_dict, keys, id_col, name_col)
            
            student_id = row_dict.get(id_col, '-')
            student_name = row_dict.get(name_col, '-')
            
            print(f"\nStudent ID: {student_id}")
            print(f"Name: {student_name}")
            print(f"  Present Days: {present}")
            print(f"  Absent Days: {absent}")
            print(f"  Attendance %: {pct}%")
            print(f"  Total Attendance Days: {present + absent}")
    else:
        print("\n✗ Daily format NOT detected")
        print("Document appears to be monthly format")
